"""Idempotent seed script for M5b development and CI.

Upserts data keyed on natural identifiers (codes, emails) — never on UUIDs.
All personal data is synthetic. Real names / emails / IDs are NEVER hardcoded
here; see CLAUDE.md seed-data rules.

Run:
    uv run python scripts/seed.py

Safe to run multiple times. Second run shows 0 rows inserted for stable
entities and 1 (upserted) for users (password re-hash on every run).
"""

from datetime import UTC, date, datetime

import structlog
from faker import Faker
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlmodel import Session, create_engine, select

import sqlalchemy as sa

from durgam.config import settings
from durgam.logging import configure_logging
from durgam.models.campus import Campus
from durgam.models.centre import CentreOfExcellence
from durgam.models.config_anchors import (
    AcademicYear,
    CalendarEntry,
    ClassTimingsConfig,
    Designation,
    DocumentTemplate,
    Holiday,
    PurchaseCommitteeTemplate,
    PurchaseProcedureRule,
    RoleEmail,
    StudentCategoryCount,
    WorkingDaysConfig,
)
from durgam.models.announcement import (
    AnnouncementCategory,
    AnnouncementComposerConfig,
    AudienceGroup,
)
from durgam.models.crosscutting import ApprovalProcess, FileAsset
from durgam.models.leave import LeaveCreditPolicy
from durgam.models.course import Course
from durgam.models.department import (
    Department,
    DepartmentCampus,
    SubDepartment,
    SubDepartmentCampus,
)
from durgam.models.identity import Permission, Role, RolePermission, User, UserRole
from durgam.models.program import (
    Program,
    ProgramExitLevel,
    ProgramOutcome,
    ProgramRegulation,
    ProgramScheme,
    ProgramSchemeCourse,
    ProgramSpecialisation,
)
from durgam.models.school import School
from durgam.models.faculty import Faculty
from durgam.models.vision_mission import (
    DepartmentMission,
    DepartmentVisionMission,
    UniversityMission,
    UniversityVisionMission,
)
from durgam.services.password import hash_password

configure_logging(debug=True)
log = structlog.get_logger(__name__)

fake = Faker(locale="en_IN")
fake.seed_instance(42)


def _exec_insert(session: Session, stmt: object) -> int:
    """Execute an INSERT...ON CONFLICT DO NOTHING and return rows actually inserted.

    Uses RETURNING to accurately count insertions (psycopg3 rowcount for
    ON CONFLICT DO NOTHING is unreliable without RETURNING).
    """
    returning = stmt.returning(sa.literal(1).label("inserted"))  # type: ignore[attr-defined]
    result = session.execute(returning)
    return len(result.fetchall())


def seed(session: Session) -> dict[str, int]:
    counts: dict[str, int] = {}

    # Dev-only: clear pre-M6a audit rows that have NULL diff_json.
    # Production never runs seed; this keeps dev/CI audit data clean for M6b UI.
    session.exec(sa.text("TRUNCATE audit_logs RESTART IDENTITY"))

    # ── AcademicYear ──────────────────────────────────────────────────────────
    ay_stmt = (
        pg_insert(AcademicYear)
        .values(
            code="2025-26",
            starts_on=date(2025, 7, 1),
            ends_on=date(2026, 4, 30),
            is_locked=False,
        )
        .on_conflict_do_nothing(constraint="uq_academic_years_code")
    )
    counts["academic_years"] = _exec_insert(session, ay_stmt)
    ay = session.exec(select(AcademicYear).where(AcademicYear.code == "2025-26")).one()

    # Second AY: 2024-25 (locked — represents a completed year; satisfies M4 gate)
    ay_prev_stmt = (
        pg_insert(AcademicYear)
        .values(
            code="2024-25",
            starts_on=date(2024, 7, 1),
            ends_on=date(2025, 4, 30),
            is_locked=True,
        )
        .on_conflict_do_nothing(constraint="uq_academic_years_code")
    )
    counts["academic_years"] += _exec_insert(session, ay_prev_stmt)
    ay_prev = session.exec(select(AcademicYear).where(AcademicYear.code == "2024-25")).one()

    # ── Roles ─────────────────────────────────────────────────────────────────
    # Role levels reflect organisational hierarchy (OQ-M4-4):
    # SYSTEM_ADMIN=100 · VC family=90-85 · REGISTRAR/FINANCE/COE=80 · DIRECTOR/CPC_CHAIR=75
    # · REGISTRAR sub=77-73 · DEPUTY_DIRECTOR=72 · IQAC=71 · DEAN_*=70
    # · DIRECTOR_OFFICE=69 · HR_HEAD=60 · HOD family=50-42 · HR_OFFICE=45
    # · LIBRARIAN etc=40 · FACULTY=30 · STUDENT=10 · BASIC_USER=1
    # on_conflict_do_update so re-seeding repairs any level drift (e.g. DEAN 50→70).
    roles_data = [
        # Technical admin (cross-cutting; not in org hierarchy)
        {"code": "SYSTEM_ADMIN",         "name": "System Administrator",    "level": 100},
        # Registrar family
        {"code": "REGISTRAR",            "name": "Registrar",               "level": 80},
        {"code": "DEPUTY_REGISTRAR",     "name": "Deputy Registrar",        "level": 77},
        {"code": "REGISTRAR_OFFICE",     "name": "Registrar Office",        "level": 73},
        # Controller of Examinations (university level; M8)
        {"code": "CONTROLLER_OF_EXAMINATIONS", "name": "Controller of Examinations", "level": 80},
        # Director family (campus-level; M4)
        {"code": "DIRECTOR",             "name": "Director",                "level": 75},
        {"code": "DEPUTY_DIRECTOR",      "name": "Deputy Director",         "level": 72},
        {"code": "DIRECTOR_OFFICE",      "name": "Director Office",         "level": 69},
        # IQAC (M4 + M10)
        {"code": "IQAC_COORDINATOR",     "name": "IQAC Coordinator",        "level": 71},
        {"code": "IQAC_OFFICE",          "name": "IQAC Office",             "level": 64},
        # Dean (school-scoped via UserRole.scope_type='school')
        {"code": "DEAN",                 "name": "Dean",                    "level": 70},
        {"code": "DEAN_STUDENT_WELFARE", "name": "Dean of Student Welfare",                          "level": 70},
        {"code": "DEAN_STUDENT_WELFARE_OFFICE", "name": "Dean of Student Welfare Office",              "level": 69},
        # Academic affairs (M5b — assignment config ownership)
        {"code": "DEAN_ACADEMIC_AFFAIRS",       "name": "Dean of Academic Affairs",                    "level": 70},
        {"code": "DEAN_ACADEMIC_AFFAIRS_OFFICE", "name": "Dean of Academic Affairs Office",            "level": 69},
        # HR family (M8 — late attendance tracking + leave admin)
        {"code": "HR_HEAD",              "name": "HR Head",                            "level": 60},
        {"code": "HR_OFFICE",            "name": "HR Office",                          "level": 45},
        # HoD family
        {"code": "HOD",                  "name": "Head of Department",                 "level": 50},
        {"code": "AHOD",                 "name": "Associate Head of Department",       "level": 45},
        {"code": "HOD_OFFICE",           "name": "Head of Department Office",          "level": 42},
        # Vice-Chancellor family (M5b — approver/channel roles for purchase config)
        {"code": "VC",                   "name": "Vice-Chancellor",                    "level": 90},
        {"code": "VC_OFFICE",            "name": "Vice-Chancellor's Office",           "level": 85},
        # Finance (M5b — owns purchase procedure rules + committee templates)
        {"code": "FINANCE_OFFICER",      "name": "Finance Officer",                    "level": 80},
        # CPC (M5b — Central Purchase Committee)
        {"code": "CPC_CHAIRPERSON",      "name": "Central Purchase Committee Chairperson", "level": 75},
        # Faculty designation roles (M8 — used in leave sanctioning matrix pre-M10 Faculty)
        # v1: users hold PROFESSOR or ASSOC_PROFESSOR role in addition to FACULTY when applicable.
        # Lecturer-tier faculty hold only FACULTY (routes to Director per §22.I.iii).
        {"code": "PROFESSOR",            "name": "Professor",                          "level": 75},
        {"code": "ASSOC_PROFESSOR",      "name": "Associate Professor",                "level": 73},
        # Faculty (dept-scoped via UserRole; M10 Faculty model deferred)
        {"code": "FACULTY",              "name": "Faculty",                            "level": 30},
        # Library, placements, centres
        {"code": "LIBRARIAN",            "name": "Librarian",                          "level": 40},
        {"code": "PLACEMENT_OFFICER",    "name": "Placement Officer",                  "level": 40},
        {"code": "CESRC_COORDINATOR",    "name": "CESRC Coordinator",                  "level": 40},
        {"code": "CENTRE_COORDINATOR",   "name": "Centre of Excellence Coordinator",   "level": 40},
        # Students and base
        {"code": "STUDENT",              "name": "Student",                            "level": 10},
        {"code": "BASIC_USER",           "name": "Basic User",                         "level": 1},
    ]
    role_counts = 0
    for r in roles_data:
        result = session.execute(
            pg_insert(Role)
            .values(**r)
            .on_conflict_do_update(
                constraint="uq_roles_code",
                set_={"name": r["name"], "level": r["level"]},
            )
            .returning(sa.literal(1).label("x"))
        )
        role_counts += len(result.fetchall())
    counts["roles"] = role_counts

    roles = {
        r.code: r
        for r in session.exec(select(Role).where(Role.is_deleted == False)).all()  # noqa: E712
    }

    # ── Permissions ───────────────────────────────────────────────────────────
    # Permissions are seed-only; no create form exists in the UI (project policy).
    # Existing M2 triples are included so a fresh DB seeds everything at once.
    perms_data = [
        # System (M2)
        {"resource": "system",           "action": "read",      "scope": "*"},
        {"resource": "system",           "action": "write",     "scope": "*"},
        {"resource": "system",           "action": "configure", "scope": "*"},
        # User management (M2)
        {"resource": "user",             "action": "read",      "scope": "*"},
        {"resource": "user",             "action": "write",     "scope": "*"},
        {"resource": "user",             "action": "delete",    "scope": "*"},
        # Role management (M2)
        {"resource": "role",             "action": "read",      "scope": "*"},
        {"resource": "role",             "action": "write",     "scope": "*"},
        {"resource": "role",             "action": "delete",    "scope": "*"},
        # Permission (M2)
        {"resource": "permission",       "action": "read",      "scope": "*"},
        # Academic year (M0)
        {"resource": "academic_year",    "action": "read",      "scope": "*"},
        {"resource": "academic_year",    "action": "write",     "scope": "*"},
        # Department — scoped variants (M2)
        {"resource": "department",       "action": "read",      "scope": "*"},
        {"resource": "department",       "action": "read",      "scope": "campus"},
        {"resource": "department",       "action": "read",      "scope": "school"},
        {"resource": "department",       "action": "read",      "scope": "department"},
        {"resource": "department",       "action": "write",     "scope": "department"},
        # Leave request (M2 placeholder — dept-scoped; extended in M8)
        {"resource": "leave_request",    "action": "read",      "scope": "department"},
        {"resource": "leave_request",    "action": "approve",   "scope": "department"},
        # Leave request — M8 full set
        {"resource": "leave_request",    "action": "create",    "scope": "*"},
        {"resource": "leave_request",    "action": "read",      "scope": "own"},
        {"resource": "leave_request",    "action": "read",      "scope": "*"},
        {"resource": "leave_request",    "action": "approve",   "scope": "*"},
        {"resource": "leave_request",    "action": "withdraw",  "scope": "own"},
        {"resource": "leave_request",    "action": "cancel",    "scope": "*"},
        # Leave balance (M8)
        {"resource": "leave_balance",    "action": "read",      "scope": "own"},
        {"resource": "leave_balance",    "action": "read",      "scope": "*"},
        {"resource": "leave_balance",    "action": "write",     "scope": "*"},
        # Leave sanction rule (M8 — sys admin configures the sanctioning matrix)
        {"resource": "leave_sanction_rule", "action": "configure", "scope": "*"},
        {"resource": "leave_sanction_rule", "action": "read",      "scope": "*"},
        # Late attendance marker (M8 — HR admin logs markers pre-attendance module)
        {"resource": "late_attendance",  "action": "write",     "scope": "*"},
        {"resource": "late_attendance",  "action": "read",      "scope": "*"},
        {"resource": "late_attendance",  "action": "read",      "scope": "department"},
        # CL credit policy (M8.1 TD-036)
        {"resource": "leave_credit_policy", "action": "read",      "scope": "*"},
        {"resource": "leave_credit_policy", "action": "configure", "scope": "*"},
        # Leave balance import (M8.1 E-016)
        {"resource": "leave_balance_import", "action": "write",     "scope": "*"},
        # Leave balance admin (M8.1 E-022)
        {"resource": "leave_balance_admin",  "action": "read",      "scope": "*"},
        {"resource": "leave_balance_admin",  "action": "write",     "scope": "*"},
        # Leave request admin (M8.1 E-022 Phase 8)
        {"resource": "leave_request_admin",  "action": "read",      "scope": "*"},
        {"resource": "leave_request_admin",  "action": "write",     "scope": "*"},
        # Audit log (M2 placeholder)
        {"resource": "audit_log",        "action": "read",      "scope": "*"},
        # ── M3 new triples ────────────────────────────────────────────────────
        # Campus
        {"resource": "campus",                     "action": "read",      "scope": "*"},
        {"resource": "campus",                     "action": "write",     "scope": "*"},
        {"resource": "campus",                     "action": "delete",    "scope": "*"},
        # School
        {"resource": "school",                     "action": "read",      "scope": "*"},
        {"resource": "school",                     "action": "write",     "scope": "*"},
        {"resource": "school",                     "action": "delete",    "scope": "*"},
        # Department write/delete at * scope (M3 adds global write)
        {"resource": "department",                 "action": "write",     "scope": "*"},
        {"resource": "department",                 "action": "delete",    "scope": "*"},
        # SubDepartment
        {"resource": "subdepartment",              "action": "read",      "scope": "*"},
        {"resource": "subdepartment",              "action": "write",     "scope": "*"},
        # Centre
        {"resource": "centre",                     "action": "read",      "scope": "*"},
        {"resource": "centre",                     "action": "write",     "scope": "*"},
        {"resource": "centre",                     "action": "delete",    "scope": "*"},
        # Program
        {"resource": "program",                    "action": "read",      "scope": "*"},
        {"resource": "program",                    "action": "write",     "scope": "*"},
        # Course
        {"resource": "course",                     "action": "read",      "scope": "*"},
        {"resource": "course",                     "action": "write",     "scope": "*"},
        {"resource": "course",                     "action": "delete",    "scope": "*"},
        # University vision/mission (E-001)
        {"resource": "university_vision_mission",  "action": "read",      "scope": "*"},
        {"resource": "university_vision_mission",  "action": "write",     "scope": "*"},
        # Department vision/mission (E-001)
        {"resource": "department_vision_mission",  "action": "read",      "scope": "*"},
        {"resource": "department_vision_mission",  "action": "write",     "scope": "department"},
        # Class timings config
        {"resource": "class_timings_config",       "action": "read",      "scope": "*"},
        {"resource": "class_timings_config",       "action": "configure", "scope": "*"},
        # Working days config
        {"resource": "working_days_config",        "action": "read",      "scope": "*"},
        {"resource": "working_days_config",        "action": "configure", "scope": "*"},
        # ── M4 new triples ────────────────────────────────────────────────────
        # Academic year — configure (lock master, manage AY lifecycle)
        {"resource": "academic_year",              "action": "configure", "scope": "*"},
        # Calendar entry
        {"resource": "calendar_entry",             "action": "read",      "scope": "*"},
        {"resource": "calendar_entry",             "action": "write",     "scope": "*"},
        # Holiday
        {"resource": "holiday",                    "action": "read",      "scope": "*"},
        {"resource": "holiday",                    "action": "write",     "scope": "*"},
        {"resource": "holiday",                    "action": "delete",    "scope": "*"},
        # Student category count
        {"resource": "student_category_count",     "action": "read",      "scope": "*"},
        {"resource": "student_category_count",     "action": "write",     "scope": "*"},
        # M5a — Role email (not public-read; Registrar family + SysAdmin only)
        {"resource": "role_email",                 "action": "read",      "scope": "*"},
        {"resource": "role_email",                 "action": "write",     "scope": "*"},
        {"resource": "role_email",                 "action": "delete",    "scope": "*"},
        # M5a — File asset download (public-read for authenticated users)
        {"resource": "file_asset",                 "action": "read",      "scope": "*"},
        # M5a — Letterhead asset (Registrar family + SysAdmin only)
        {"resource": "letterhead_asset",           "action": "read",      "scope": "*"},
        {"resource": "letterhead_asset",           "action": "write",     "scope": "*"},
        {"resource": "letterhead_asset",           "action": "delete",    "scope": "*"},
        # M5a — Template asset (IQAC + SysAdmin only; NOT in _PUBLIC_READ)
        {"resource": "template_asset",             "action": "read",      "scope": "*"},
        {"resource": "template_asset",             "action": "write",     "scope": "*"},
        {"resource": "template_asset",             "action": "delete",    "scope": "*"},
        # ── M5b new triples ────────────────────────────────────────────────────
        # Mental health counsellor (Dean SW family + SysAdmin)
        {"resource": "mental_health_counsellor",   "action": "read",      "scope": "*"},
        {"resource": "mental_health_counsellor",   "action": "write",     "scope": "*"},
        {"resource": "mental_health_counsellor",   "action": "delete",    "scope": "*"},
        # Faculty mentor assignment (Dean SW family + SysAdmin)
        {"resource": "faculty_mentor_assignment",  "action": "read",      "scope": "*"},
        {"resource": "faculty_mentor_assignment",  "action": "write",     "scope": "*"},
        {"resource": "faculty_mentor_assignment",  "action": "delete",    "scope": "*"},
        # Class teacher assignment (Dean Academic Affairs family + HOD + SysAdmin)
        {"resource": "class_teacher_assignment",   "action": "read",      "scope": "*"},
        {"resource": "class_teacher_assignment",   "action": "write",     "scope": "*"},
        {"resource": "class_teacher_assignment",   "action": "delete",    "scope": "*"},
        # Class coordinator assignment (Dean Academic Affairs family + HOD + SysAdmin)
        {"resource": "class_coordinator_assignment", "action": "read",    "scope": "*"},
        {"resource": "class_coordinator_assignment", "action": "write",   "scope": "*"},
        {"resource": "class_coordinator_assignment", "action": "delete",  "scope": "*"},
        # Non-regular faculty (HoD family read/write/delete + SysAdmin; approve = SysAdmin only)
        {"resource": "non_regular_faculty",          "action": "read",    "scope": "*"},
        {"resource": "non_regular_faculty",          "action": "write",   "scope": "*"},
        {"resource": "non_regular_faculty",          "action": "delete",  "scope": "*"},
        {"resource": "non_regular_faculty",          "action": "approve", "scope": "*"},
        # Non-owned course (Director family + DAA family + SysAdmin; no department_id)
        {"resource": "non_owned_course",             "action": "read",    "scope": "*"},
        {"resource": "non_owned_course",             "action": "write",   "scope": "*"},
        {"resource": "non_owned_course",             "action": "delete",  "scope": "*"},
        # UG timetable (Director family + SysAdmin only)
        {"resource": "ug_timetable",                 "action": "read",    "scope": "*"},
        {"resource": "ug_timetable",                 "action": "write",   "scope": "*"},
        {"resource": "ug_timetable",                 "action": "delete",  "scope": "*"},
        # ── M5b Session 7: Purchase policy & approval config ──────────────────
        # Purchase procedure rule (Finance Officer only — NOT in _PUBLIC_READ; E-007)
        {"resource": "purchase_procedure_rule",      "action": "read",    "scope": "*"},
        {"resource": "purchase_procedure_rule",      "action": "write",   "scope": "*"},
        {"resource": "purchase_procedure_rule",      "action": "delete",  "scope": "*"},
        # Purchase committee template (Finance Officer only — NOT in _PUBLIC_READ; E-007)
        {"resource": "purchase_committee_template",  "action": "read",    "scope": "*"},
        {"resource": "purchase_committee_template",  "action": "write",   "scope": "*"},
        {"resource": "purchase_committee_template",  "action": "delete",  "scope": "*"},
        # Approval process (SysAdmin only via global — NOT in _PUBLIC_READ)
        {"resource": "approval_process",             "action": "read",    "scope": "*"},
        {"resource": "approval_process",             "action": "write",   "scope": "*"},
        {"resource": "approval_process",             "action": "delete",  "scope": "*"},
        # Approval request — approver nav visibility (M7)
        {"resource": "approval_request",             "action": "approve", "scope": "*"},
        # Designation vocabulary (Finance Officer + SysAdmin)
        {"resource": "designation",                  "action": "read",    "scope": "*"},
        {"resource": "designation",                  "action": "write",   "scope": "*"},
        {"resource": "designation",                  "action": "delete",  "scope": "*"},
        # M5b-R3 V2 — per-entity bulk-import permissions
        {"resource": "program_import",               "action": "write",   "scope": "*"},
        {"resource": "course_import",                "action": "write",   "scope": "*"},
        # M9 — Announcement Module permissions
        {"resource": "announcement",                 "action": "create",    "scope": "*"},
        {"resource": "announcement",                 "action": "read",      "scope": "*"},
        {"resource": "announcement",                 "action": "update",    "scope": "own"},
        {"resource": "announcement",                 "action": "soft_delete", "scope": "own"},
        {"resource": "announcement_composer_config", "action": "read",      "scope": "*"},
        {"resource": "announcement_composer_config", "action": "configure", "scope": "*"},
        {"resource": "announcement_category",        "action": "read",      "scope": "*"},
        {"resource": "announcement_category",        "action": "configure", "scope": "*"},
        {"resource": "audience_group",               "action": "read",      "scope": "*"},
        {"resource": "audience_group",               "action": "configure", "scope": "*"},
        # ── M10 Faculty Module permissions ────────────────────────────────────
        # Faculty profile — all authenticated can read directory; own/admin write split
        {"resource": "faculty",           "action": "read",        "scope": "*"},
        {"resource": "faculty",           "action": "write",       "scope": "own"},
        {"resource": "faculty",           "action": "write",       "scope": "*"},   # admin-tier (Registrar/HR_HEAD)
        # Sensitive PII fields — Registrar + IQAC family only (RFP §9.7)
        {"resource": "faculty_sensitive", "action": "read",        "scope": "*"},
        # Faculty-uploaded documents — Registrar + IQAC family only
        {"resource": "faculty_document",  "action": "read",        "scope": "*"},
        # Faculty requests — own create/read; * read for Registrar audit
        {"resource": "faculty_request",   "action": "create",      "scope": "own"},
        {"resource": "faculty_request",   "action": "read",        "scope": "own"},
        {"resource": "faculty_request",   "action": "read",        "scope": "*"},
        # Faculty workload — all read (for admin); own write (self-entry)
        {"resource": "faculty_workload",  "action": "read",        "scope": "*"},
        {"resource": "faculty_workload",  "action": "write",       "scope": "own"},
        # Faculty bulk import — Registrar + HR_HEAD
        {"resource": "faculty",           "action": "bulk_import", "scope": "*"},
        # Designation configure — SYSTEM_ADMIN only (vocabulary-level change)
        {"resource": "designation",       "action": "configure",   "scope": "*"},
        # ApprovalProcess configure — SYSTEM_ADMIN only (extended for OR-set channels at M10)
        {"resource": "approval_process",  "action": "configure",   "scope": "*"},
    ]
    perm_inserted = 0
    for p in perms_data:
        perm_inserted += _exec_insert(
            session,
            pg_insert(Permission)
            .values(**p)
            .on_conflict_do_nothing(constraint="uq_permissions_resource_action_scope"),
        )
    counts["permissions"] = perm_inserted

    perms = {
        (p.resource, p.action, p.scope): p
        for p in session.exec(
            select(Permission).where(Permission.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── RolePermissions ───────────────────────────────────────────────────────
    # SYSTEM_ADMIN: all permissions.
    rp_inserted = 0
    for perm in perms.values():
        rp_inserted += _exec_insert(
            session,
            pg_insert(RolePermission)
            .values(role_id=roles["SYSTEM_ADMIN"].id, permission_id=perm.id)
            .on_conflict_do_nothing(),
        )

    # Permissions shared by all read-capable roles (everything M3 added as read:*)
    _PUBLIC_READ = [
        ("campus",                    "read", "*"),
        ("school",                    "read", "*"),
        ("department",                "read", "*"),
        ("department",                "read", "campus"),
        ("department",                "read", "school"),
        ("department",                "read", "department"),
        ("subdepartment",             "read", "*"),
        ("centre",                    "read", "*"),
        ("program",                   "read", "*"),
        ("course",                    "read", "*"),
        ("university_vision_mission", "read", "*"),
        ("department_vision_mission", "read", "*"),
        ("class_timings_config",      "read", "*"),
        ("working_days_config",       "read", "*"),
        ("academic_year",             "read", "*"),
        ("holiday",                   "read", "*"),
        # M5a — file download permission for all authenticated users
        ("file_asset",                "read", "*"),
        # M5b — non-owned courses + UG timetable (scheduling info all users see)
        ("non_owned_course",          "read", "*"),
        ("ug_timetable",              "read", "*"),
        # M5b — class coordinator list viewable by all (B1: coordinator is a student)
        ("class_coordinator_assignment", "read", "*"),
        # M9 — Announcement Module reads (broad transparency: every role can see announcements,
        # the composer roster, categories, and audience groups)
        ("announcement",                 "read", "*"),
        ("announcement_composer_config", "read", "*"),
        ("announcement_category",        "read", "*"),
        ("audience_group",               "read", "*"),
    ]

    # M8 — permissions for all employees who can submit leave requests.
    # Excludes STUDENT and BASIC_USER (no leave entitlement).
    _LEAVE_REQUESTOR = [
        ("leave_request", "create",   "*"),
        ("leave_request", "read",     "own"),
        ("leave_request", "withdraw", "own"),
        ("leave_balance", "read",     "own"),
    ]

    _REGISTRAR_SPECIFIC = [
        ("university_vision_mission",  "write",     "*"),
        ("class_timings_config",       "configure", "*"),
        ("working_days_config",        "configure", "*"),
        # department:write:* intentionally NOT included — §9.3: only SYSTEM_ADMIN
        # manages department structure. Registrar has read access via _PUBLIC_READ.
        # M4 — Registrar owns AY, holiday, student category, and calendar master entries
        ("academic_year",              "write",     "*"),
        ("academic_year",              "configure", "*"),
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("holiday",                    "write",     "*"),
        ("holiday",                    "delete",    "*"),
        ("student_category_count",     "read",      "*"),
        ("student_category_count",     "write",     "*"),
        # M5a — role email management (not in _PUBLIC_READ — internal service lookup only)
        ("role_email",                 "read",      "*"),
        ("role_email",                 "write",     "*"),
        ("role_email",                 "delete",    "*"),
        # M5a — letterhead asset management (Registrar family + SysAdmin)
        ("letterhead_asset",           "read",      "*"),
        ("letterhead_asset",           "write",     "*"),
        ("letterhead_asset",           "delete",    "*"),
        # M5b-R3 V1 — Registrar family manages programs
        ("program",                    "write",     "*"),
        ("program",                    "delete",    "*"),
        # M5b-R3 V2 — Registrar family can bulk-import programs
        ("program_import",             "write",     "*"),
        # M5b — non-regular faculty approval (Registrar family is also institutional approver)
        ("non_regular_faculty",        "approve",   "*"),
        # M8 — Registrar family can view + configure the leave sanction matrix
        ("leave_sanction_rule",        "configure", "*"),
        ("leave_sanction_rule",        "read",      "*"),
    ]

    _HOD_SPECIFIC = [
        ("department_vision_mission",  "write",     "department"),
        ("leave_request",              "read",      "department"),
        ("leave_request",              "approve",   "department"),
        # M4 — HOD can read calendar and create department entries
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
        # M5b — non-regular faculty (read/write/delete but NOT approve)
        ("non_regular_faculty",        "read",      "*"),
        ("non_regular_faculty",        "write",     "*"),
        ("non_regular_faculty",        "delete",    "*"),
        # M5b-R3 V1 — HOD manages courses for their department
        ("course",                     "write",     "*"),
        ("course",                     "delete",    "*"),
        # M5b-R3 V2 — HOD can bulk-import courses
        ("course_import",              "write",     "*"),
    ]

    _DEAN_SPECIFIC = [
        ("department",                 "read",      "school"),
        ("leave_request",              "read",      "department"),
        ("leave_request",              "approve",   "department"),
        # M4 — Deans can read/write calendar (Phase 3 generic types) + read student category
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
    ]

    # M4 — Director family can read/write calendar, read student category
    # M5b — Director family owns counsellor roster + faculty mentor assignments
    _DIRECTOR_SPECIFIC = [
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
        ("mental_health_counsellor",   "read",      "*"),
        ("mental_health_counsellor",   "write",     "*"),
        ("mental_health_counsellor",   "delete",    "*"),
        ("faculty_mentor_assignment",  "read",      "*"),
        ("faculty_mentor_assignment",  "write",     "*"),
        ("faculty_mentor_assignment",  "delete",    "*"),
        # M5b — non-owned courses (Director + DAA) + UG timetable (Director only)
        ("non_owned_course",           "read",      "*"),
        ("non_owned_course",           "write",     "*"),
        ("non_owned_course",           "delete",    "*"),
        ("ug_timetable",               "read",      "*"),
        ("ug_timetable",               "write",     "*"),
        ("ug_timetable",               "delete",    "*"),
        # M5b — non-regular faculty approval (Director is institutional approver §7.1)
        ("non_regular_faculty",        "approve",   "*"),
        # M8 — Director family logs late-attendance markers for their campus (scope:* for v1;
        # campus-scoped filtering deferred to post-M10 once Faculty/Campus assignment hardens)
        ("late_attendance",            "write",     "*"),
        ("late_attendance",            "read",      "*"),
    ]

    # M4 — IQAC can read/write calendar, read student category
    # M5a — IQAC owns template assets
    _IQAC_SPECIFIC = [
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
        ("template_asset",             "read",      "*"),
        ("template_asset",             "write",     "*"),
        ("template_asset",             "delete",    "*"),
    ]

    # M4 — Dean of Student Welfare can read/write calendar, read student category
    # M5b — Dean SW owns counsellor + faculty mentor assignment config
    _DEAN_SW_SPECIFIC = [
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
        ("mental_health_counsellor",   "read",      "*"),
        ("mental_health_counsellor",   "write",     "*"),
        ("mental_health_counsellor",   "delete",    "*"),
        ("faculty_mentor_assignment",  "read",      "*"),
        ("faculty_mentor_assignment",  "write",     "*"),
        ("faculty_mentor_assignment",  "delete",    "*"),
    ]

    # M5b — Dean Academic Affairs: calendar, student categories, non-owned courses
    # Class teacher/coordinator read-only (write is HoD-managed, not DAA)
    _DEAN_AA_SPECIFIC = [
        ("calendar_entry",             "read",      "*"),
        ("calendar_entry",             "write",     "*"),
        ("student_category_count",     "read",      "*"),
        ("class_teacher_assignment",   "read",      "*"),
        ("class_coordinator_assignment", "read",    "*"),
        # M5b — non-owned courses (DAA family; NOT ug_timetable — Director only)
        ("non_owned_course",           "read",      "*"),
        ("non_owned_course",           "write",     "*"),
        ("non_owned_course",           "delete",    "*"),
    ]

    # M5b — Finance Officer owns purchase procedure rules + committee templates +
    # designation vocabulary (E-007: "not viewable by others" → NOT in _PUBLIC_READ)
    _FINANCE_SPECIFIC = [
        ("purchase_procedure_rule",     "read",   "*"),
        ("purchase_procedure_rule",     "write",  "*"),
        ("purchase_procedure_rule",     "delete", "*"),
        ("purchase_committee_template", "read",   "*"),
        ("purchase_committee_template", "write",  "*"),
        ("purchase_committee_template", "delete", "*"),
        ("designation",                 "read",   "*"),
        ("designation",                 "write",  "*"),
        ("designation",                 "delete", "*"),
        # M7 — approval request approver (channel role in CPC_FUND_RELEASE)
        ("approval_request",            "approve", "*"),
    ]

    # M8 — Controller of Examinations: same approval scope as VC/Registrar for leave
    _CONTROLLER_OF_EXAMINATIONS_SPECIFIC = [
        ("leave_request",              "read",      "department"),
        ("leave_request",              "approve",   "*"),
        ("approval_request",           "approve",   "*"),
    ]

    # M8 — HR Head: leave admin (read all requests + balances + write late-attendance)
    _HR_HEAD_SPECIFIC = [
        ("leave_request",              "read",      "*"),
        ("leave_balance",              "read",      "*"),
        ("late_attendance",            "write",     "*"),
        ("late_attendance",            "read",      "department"),
    ]

    # M8 — HR Office: read leave requests + balances (no write on late attendance)
    _HR_OFFICE_SPECIFIC = [
        ("leave_request",              "read",      "*"),
        ("leave_balance",              "read",      "*"),
        ("late_attendance",            "read",      "department"),
    ]

    # M8.1 TD-036 — CL credit policy admin (read + configure)
    _CL_CREDIT_POLICY_SPECIFIC = [
        ("leave_credit_policy",        "read",      "*"),
        ("leave_credit_policy",        "configure", "*"),
    ]

    # M8.1 E-016 — leave balance import (REGISTRAR family + DIRECTOR family)
    _LEAVE_BALANCE_IMPORT = [
        ("leave_balance_import",       "write",     "*"),
    ]

    # M8.1 E-022 — leave balance admin edit
    _LEAVE_BALANCE_ADMIN = [
        ("leave_balance_admin",        "read",      "*"),
        ("leave_balance_admin",        "write",     "*"),
    ]

    # M8.1 E-022 Phase 8 — leave request admin edit
    _LEAVE_REQUEST_ADMIN = [
        ("leave_request_admin",        "read",      "*"),
        ("leave_request_admin",        "write",     "*"),
    ]

    # M9 — Announcement Module
    _ANNOUNCEMENT_COMPOSER = [
        ("announcement",   "create",      "*"),
        ("announcement",   "update",      "own"),
        ("announcement",   "soft_delete", "own"),
    ]

    # Registrar-tier operational config (categories + audience groups)
    _ANNOUNCEMENT_REGISTRAR_CONFIG = [
        ("announcement_category", "configure", "*"),
        ("audience_group",        "configure", "*"),
    ]

    # Sys-admin-only configuration of the composer roster itself.
    # SYSTEM_ADMIN receives this automatically via the perms.values() loop above;
    # this list is defined here for clarity but not added to role_perm_map.
    _ANNOUNCEMENT_SYS_ADMIN_CONFIG = [
        ("announcement_composer_config", "configure", "*"),
    ]

    # M10 — Faculty Module (D-011)
    # Base set for all regular-teaching employees: directory read + self-edit + own requests.
    _FACULTY_OWN = [
        ("faculty",          "read",        "*"),    # all authenticated can see faculty directory
        ("faculty",          "write",       "own"),  # self-edit own profile
        ("faculty_request",  "create",      "own"),  # raise faculty requests
        ("faculty_request",  "read",        "own"),  # track own requests
        ("faculty_workload", "write",       "own"),  # enter own workload
        ("faculty_workload", "read",        "*"),    # view workload (own included in *)
    ]

    # Admin-tier write: Registrar/HR_HEAD can edit any faculty record + bulk import + audit
    _FACULTY_ADMIN = [
        ("faculty",          "write",       "*"),    # admin-level write any faculty record
        ("faculty_request",  "read",        "*"),    # audit: view all faculty requests
        ("faculty",          "bulk_import", "*"),    # bulk import
    ]

    # Sensitive PII read: Registrar + IQAC family only (RFP §9.7)
    _FACULTY_SENSITIVE_READ = [
        ("faculty_sensitive", "read",       "*"),    # PAN/Aadhaar/sensitive fields
        ("faculty_document",  "read",       "*"),    # uploaded faculty documents
    ]

    role_perm_map: dict[str, list[tuple[str, str, str]]] = {
        "REGISTRAR":            _PUBLIC_READ + _LEAVE_REQUESTOR + _REGISTRAR_SPECIFIC + _CL_CREDIT_POLICY_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _ANNOUNCEMENT_COMPOSER + _ANNOUNCEMENT_REGISTRAR_CONFIG + _FACULTY_OWN + _FACULTY_ADMIN + _FACULTY_SENSITIVE_READ + [
            ("approval_request",           "approve",   "*"),
        ],
        "DEPUTY_REGISTRAR":     _PUBLIC_READ + _LEAVE_REQUESTOR + _REGISTRAR_SPECIFIC + _CL_CREDIT_POLICY_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _FACULTY_OWN + _FACULTY_ADMIN + _FACULTY_SENSITIVE_READ + [
            ("approval_request",           "approve",   "*"),
        ],
        "REGISTRAR_OFFICE":     _PUBLIC_READ + _LEAVE_REQUESTOR + _REGISTRAR_SPECIFIC + _CL_CREDIT_POLICY_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _ANNOUNCEMENT_COMPOSER + _ANNOUNCEMENT_REGISTRAR_CONFIG + _FACULTY_OWN + _FACULTY_ADMIN + _FACULTY_SENSITIVE_READ,
        "DIRECTOR":             _PUBLIC_READ + _LEAVE_REQUESTOR + _DIRECTOR_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "DEPUTY_DIRECTOR":      _PUBLIC_READ + _LEAVE_REQUESTOR + _DIRECTOR_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _FACULTY_OWN,
        "DIRECTOR_OFFICE":      _PUBLIC_READ + _LEAVE_REQUESTOR + _DIRECTOR_SPECIFIC + _LEAVE_BALANCE_IMPORT + _LEAVE_BALANCE_ADMIN + _LEAVE_REQUEST_ADMIN + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "IQAC_COORDINATOR":     _PUBLIC_READ + _LEAVE_REQUESTOR + _IQAC_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + _FACULTY_SENSITIVE_READ,
        # M10 — IQAC_OFFICE mirrors IQAC_COORDINATOR exactly (Q-P2.1 authority 2026-06-14)
        "IQAC_OFFICE":          _PUBLIC_READ + _LEAVE_REQUESTOR + _IQAC_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + _FACULTY_SENSITIVE_READ,
        "DEAN":                 _PUBLIC_READ + _LEAVE_REQUESTOR + _DEAN_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + [
            ("approval_request",           "approve",   "*"),
        ],
        "DEAN_STUDENT_WELFARE": _PUBLIC_READ + _LEAVE_REQUESTOR + _DEAN_SW_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "DEAN_STUDENT_WELFARE_OFFICE": _PUBLIC_READ + _LEAVE_REQUESTOR + _DEAN_SW_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "DEAN_ACADEMIC_AFFAIRS": _PUBLIC_READ + _LEAVE_REQUESTOR + _DEAN_AA_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "DEAN_ACADEMIC_AFFAIRS_OFFICE": _PUBLIC_READ + _LEAVE_REQUESTOR + _DEAN_AA_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "HOD":                  _PUBLIC_READ + _LEAVE_REQUESTOR + _HOD_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + [
            ("faculty_request",            "read",      "*"),  # dept scope enforced at handler body (M3 pattern)
            ("class_teacher_assignment",   "read",      "*"),
            ("class_teacher_assignment",   "write",     "*"),
            ("class_teacher_assignment",   "delete",    "*"),
            ("class_coordinator_assignment", "read",    "*"),
            ("class_coordinator_assignment", "write",   "*"),
            ("class_coordinator_assignment", "delete",  "*"),
        ],
        "AHOD":                 _PUBLIC_READ + _LEAVE_REQUESTOR + _HOD_SPECIFIC + _FACULTY_OWN + [
            ("faculty_request",            "read",      "*"),  # dept scope enforced at handler body (M3 pattern)
            ("class_teacher_assignment",   "read",      "*"),
            ("class_teacher_assignment",   "write",     "*"),
            ("class_teacher_assignment",   "delete",    "*"),
            ("class_coordinator_assignment", "read",    "*"),
            ("class_coordinator_assignment", "write",   "*"),
            ("class_coordinator_assignment", "delete",  "*"),
        ],
        "HOD_OFFICE":           _PUBLIC_READ + _LEAVE_REQUESTOR + [
            ("calendar_entry",             "read",      "*"),
            ("calendar_entry",             "write",     "*"),
            ("student_category_count",     "read",      "*"),
            # M5b — non-regular faculty read-only
            ("non_regular_faculty",        "read",      "*"),
            # M5b-R3 V1 — HoD Office assists with course and dept V&M management
            ("course",                     "write",     "*"),
            ("course",                     "delete",    "*"),
            ("department_vision_mission",  "write",     "department"),
            # M5b-R3 V2 — HoD Office can bulk-import courses
            ("course_import",              "write",     "*"),
            # M10 — faculty directory read (no self-edit: HOD_OFFICE is admin staff)
            ("faculty",                    "read",      "*"),
        ],
        # M5b/M7 — approver/channel roles
        "VC":                   _PUBLIC_READ + _LEAVE_REQUESTOR + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + [
            ("approval_request",           "approve",   "*"),
        ],
        "VC_OFFICE":            _PUBLIC_READ + _LEAVE_REQUESTOR + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "FINANCE_OFFICER":      _PUBLIC_READ + _LEAVE_REQUESTOR + _FINANCE_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "CPC_CHAIRPERSON":      _PUBLIC_READ + _LEAVE_REQUESTOR + [
            ("approval_request",           "approve",   "*"),
            ("faculty",                    "read",      "*"),
        ],
        # M8 — new roles
        "CONTROLLER_OF_EXAMINATIONS": _PUBLIC_READ + _LEAVE_REQUESTOR + _CONTROLLER_OF_EXAMINATIONS_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "HR_HEAD":              _PUBLIC_READ + _LEAVE_REQUESTOR + _HR_HEAD_SPECIFIC + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN + _FACULTY_ADMIN + _FACULTY_SENSITIVE_READ,
        "HR_OFFICE":            _PUBLIC_READ + _LEAVE_REQUESTOR + _HR_OFFICE_SPECIFIC + _FACULTY_OWN + _FACULTY_ADMIN + _FACULTY_SENSITIVE_READ,
        # Faculty designation roles (M8 — inherit PUBLIC_READ + LEAVE_REQUESTOR; M10 adds faculty perms)
        "PROFESSOR":            _PUBLIC_READ + _LEAVE_REQUESTOR + _FACULTY_OWN + [("approval_request", "approve", "*")],
        "ASSOC_PROFESSOR":      _PUBLIC_READ + _LEAVE_REQUESTOR + _FACULTY_OWN + [("approval_request", "approve", "*")],
        "FACULTY":              _PUBLIC_READ + _LEAVE_REQUESTOR + _FACULTY_OWN,
        "LIBRARIAN":            _PUBLIC_READ + _LEAVE_REQUESTOR + _FACULTY_OWN,
        "PLACEMENT_OFFICER":    _PUBLIC_READ + _LEAVE_REQUESTOR + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "CESRC_COORDINATOR":    _PUBLIC_READ + _LEAVE_REQUESTOR + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "CENTRE_COORDINATOR":   _PUBLIC_READ + _LEAVE_REQUESTOR + _ANNOUNCEMENT_COMPOSER + _FACULTY_OWN,
        "STUDENT":              _PUBLIC_READ,
        "BASIC_USER":           _PUBLIC_READ,
    }

    for role_code, perm_keys in role_perm_map.items():
        role = roles[role_code]
        seen: set[tuple[str, str, str]] = set()
        for key in perm_keys:
            if key in seen:
                continue
            seen.add(key)
            if key not in perms:
                continue
            rp_inserted += _exec_insert(
                session,
                pg_insert(RolePermission)
                .values(role_id=role.id, permission_id=perms[key].id)
                .on_conflict_do_nothing(),
            )

    counts["role_permissions"] = rp_inserted

    # ── Cleanup: remove incorrect REGISTRAR-department:write:* assignments ────
    # department:write:* was erroneously assigned to the Registrar family at M3
    # Session 2. Only SYSTEM_ADMIN manages department structure per §9.3.
    # This DELETE is idempotent — safe to re-run if already cleaned up.
    if ("department", "write", "*") in perms:
        dept_write_perm = perms[("department", "write", "*")]
        for role_code in ("REGISTRAR", "DEPUTY_REGISTRAR", "REGISTRAR_OFFICE"):
            if role_code in roles:
                session.execute(
                    sa.delete(RolePermission).where(
                        RolePermission.role_id == roles[role_code].id,
                        RolePermission.permission_id == dept_write_perm.id,
                    )
                )

    # ── Users ─────────────────────────────────────────────────────────────────
    # Read-only seeded fixtures (CLAUDE.md Testing rules):
    #   sys_admin / SysAdmin_Dev1!XZ      — SYSTEM_ADMIN
    #   dean_sci / DeanSci_Dev1!XZ        — DEAN scoped to SCI school
    #   firstlogin_user / FirstLogin_Dev1!XZ — STUDENT, must_change_password=True
    #   inactive_user / Inactive_Dev1!XZ  — STUDENT, is_active=False
    #   student_001 / Student_Dev1!XZ     — STUDENT
    #   registrar_user / Registrar_Dev1!XZ — REGISTRAR  (new at M3)
    #   hod_dmacs / HodDmacs_Dev1!XZ     — HOD scoped to DMACS (new at M3; scoped role added after depts)
    #   director_psn / DirectorPsn_Dev1!XZ — DIRECTOR scoped to PSN (new at M4; scoped role added after campuses)
    #   iqac_user / IqacCoord_Dev1!XZ    — IQAC_COORDINATOR unscoped (new at M4)
    #   dean_sw / DeanSW_Dev1!XZ         — DEAN_STUDENT_WELFARE unscoped (new at M4)
    #   finance_user / Finance_Dev1!XZ  — FINANCE_OFFICER unscoped (new at M5b)
    #   daa_user / DeanAA_Dev1!XZ       — DEAN_ACADEMIC_AFFAIRS unscoped (new at M5b)
    #   registrar_office_user / RegOffice_Dev1!XZ — REGISTRAR_OFFICE (M5b-R2)
    #   deputy_registrar_user / DeputyReg_Dev1!XZ — DEPUTY_REGISTRAR (M5b-R2)
    #   hod_office_dmacs / HodOffice_Dev1!XZ — HOD_OFFICE scoped to DMACS (M5b-R2)
    #   ahod_dmacs / AhodDmacs_Dev1!XZ — AHOD scoped to DMACS (M5b-R2)
    #   director_office_psn / DirOffice_Dev1!XZ — DIRECTOR_OFFICE scoped to PSN (M5b-R2)
    #   deputy_director_psn / DeputyDir_Dev1!XZ — DEPUTY_DIRECTOR scoped to PSN (M5b-R2)
    #   dsw_office_user / DSWOffice_Dev1!XZ — DEAN_STUDENT_WELFARE_OFFICE (M5b-R2)
    #   daa_office_user / DaaOffice_Dev1!XZ — DEAN_ACADEMIC_AFFAIRS_OFFICE (M5b-R2)
    #   faculty_user / Faculty_Dev1!XZ — FACULTY scoped to DMACS (M5b-R2)
    #   librarian_user / Librarian_Dev1!XZ — LIBRARIAN (M5b-R2)
    #   placement_officer_user / Placement_Dev1!XZ — PLACEMENT_OFFICER (M5b-R2)
    #   cesrc_coord_user / Cesrc_Dev1!XZ — CESRC_COORDINATOR (M5b-R2)
    #   center_coord_user / Center_Dev1!XZ — CENTRE_COORDINATOR (M5b-R2)
    #   vc_user / ViceChancellor_Dev1!XZ — VC unscoped (M8)
    users_data = [
        # M8 employment fields added: gender ('M'|'F'|'O'|None), joined_on (date|None),
        # employee_type (regular_teaching|regular_non_teaching|honorary_*|superannuated_*|visiting_fellow)
        # faculty_user is gender="F", joined_on=date(2022,6,1) — maternity-eligible fixture (≥1y service).
        {
            "email": "sys.admin@sssihl.edu.in",
            "username": "sys_admin",
            "full_name": "System Administrator",
            "role_code": "SYSTEM_ADMIN",
            "plain_password": "SysAdmin_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2015, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "dean.sci@sssihl.edu.in",
            "username": "dean_sci",
            "full_name": "Dean Sciences",
            "role_code": "DEAN",
            "plain_password": "DeanSci_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_teaching",
        },
        {
            "email": "student.001@sssihl.edu.in",
            "username": "student_001",
            "full_name": "Student One",
            "role_code": "STUDENT",
            "plain_password": "Student_Dev1!XZ",
        },
        {
            "email": "inactive.user@sssihl.edu.in",
            "username": "inactive_user",
            "full_name": "Inactive User",
            "role_code": "STUDENT",
            "plain_password": "Inactive_Dev1!XZ",
            "is_active": False,
        },
        {
            "email": "firstlogin.user@sssihl.edu.in",
            "username": "firstlogin_user",
            "full_name": "First Login User",
            "role_code": "STUDENT",
            "plain_password": "FirstLogin_Dev1!XZ",
            "must_change_password": True,
        },
        # M3 demo users
        {
            "email": "registrar@sssihl.edu.in",
            "username": "registrar_user",
            "full_name": "University Registrar",
            "role_code": "REGISTRAR",
            "plain_password": "Registrar_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "hod.dmacs@sssihl.edu.in",
            "username": "hod_dmacs",
            "full_name": "HoD Mathematics and Computer Science",
            # HOD role is scoped to DMACS — assigned after departments are seeded.
            # Only BASIC_USER is assigned here; see the scoped-roles block below.
            "role_code": "BASIC_USER",
            "plain_password": "HodDmacs_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_teaching",
        },
        # M4 demo users
        {
            "email": "director.psn@sssihl.edu.in",
            "username": "director_psn",
            "full_name": "Director Prasanthi Nilayam",
            # DIRECTOR role is campus-scoped — assigned after campuses are seeded.
            "role_code": "BASIC_USER",
            "plain_password": "DirectorPsn_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2015, 6, 1),
            "employee_type": "regular_teaching",
        },
        {
            "email": "iqac.coordinator@sssihl.edu.in",
            "username": "iqac_user",
            "full_name": "IQAC Coordinator",
            "role_code": "IQAC_COORDINATOR",
            "plain_password": "IqacCoord_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "dean.sw@sssihl.edu.in",
            "username": "dean_sw",
            "full_name": "Dean of Student Welfare",
            "role_code": "DEAN_STUDENT_WELFARE",
            "plain_password": "DeanSW_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        # M5b demo users
        {
            "email": "finance.officer@sssihl.edu.in",
            "username": "finance_user",
            "full_name": "Finance Officer",
            "role_code": "FINANCE_OFFICER",
            "plain_password": "Finance_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "dean.aa@sssihl.edu.in",
            "username": "daa_user",
            "full_name": "Dean of Academic Affairs",
            "role_code": "DEAN_ACADEMIC_AFFAIRS",
            "plain_password": "DeanAA_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        # M5b-R2 demo users (13 new)
        {
            "email": "registrar.office@sssihl.edu.in",
            "username": "registrar_office_user",
            "full_name": "Registrar Office Staff",
            "role_code": "REGISTRAR_OFFICE",
            "plain_password": "RegOffice_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "deputy.registrar@sssihl.edu.in",
            "username": "deputy_registrar_user",
            "full_name": "Deputy Registrar",
            "role_code": "DEPUTY_REGISTRAR",
            "plain_password": "DeputyReg_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "hod.office.dmacs@sssihl.edu.in",
            "username": "hod_office_dmacs",
            "full_name": "HoD Office DMACS",
            "role_code": "BASIC_USER",
            "plain_password": "HodOffice_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "ahod.dmacs@sssihl.edu.in",
            "username": "ahod_dmacs",
            "full_name": "Associate HoD Mathematics and Computer Science",
            "role_code": "BASIC_USER",
            "plain_password": "AhodDmacs_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_teaching",
        },
        {
            "email": "director.office.psn@sssihl.edu.in",
            "username": "director_office_psn",
            "full_name": "Director Office Prasanthi Nilayam",
            "role_code": "BASIC_USER",
            "plain_password": "DirOffice_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "deputy.director.psn@sssihl.edu.in",
            "username": "deputy_director_psn",
            "full_name": "Deputy Director Prasanthi Nilayam",
            "role_code": "BASIC_USER",
            "plain_password": "DeputyDir_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_teaching",
        },
        {
            "email": "dsw.office@sssihl.edu.in",
            "username": "dsw_office_user",
            "full_name": "Dean Student Welfare Office Staff",
            "role_code": "DEAN_STUDENT_WELFARE_OFFICE",
            "plain_password": "DSWOffice_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "daa.office@sssihl.edu.in",
            "username": "daa_office_user",
            "full_name": "Dean Academic Affairs Office Staff",
            "role_code": "DEAN_ACADEMIC_AFFAIRS_OFFICE",
            "plain_password": "DaaOffice_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "faculty.dmacs@sssihl.edu.in",
            "username": "faculty_user",
            "full_name": "Faculty Member DMACS",
            "role_code": "BASIC_USER",
            "plain_password": "Faculty_Dev1!XZ",
            # M8: gender="F", joined_on 2022-06-01 → ≥1y service by 2026-06-08 → ML-eligible fixture
            "gender": "F",
            "joined_on": date(2022, 6, 1),
            "employee_type": "regular_teaching",
        },
        {
            "email": "librarian@sssihl.edu.in",
            "username": "librarian_user",
            "full_name": "University Librarian",
            "role_code": "LIBRARIAN",
            "plain_password": "Librarian_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "placement@sssihl.edu.in",
            "username": "placement_officer_user",
            "full_name": "Placement Officer",
            "role_code": "PLACEMENT_OFFICER",
            "plain_password": "Placement_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "cesrc@sssihl.edu.in",
            "username": "cesrc_coord_user",
            "full_name": "CESRC Coordinator",
            "role_code": "CESRC_COORDINATOR",
            "plain_password": "Cesrc_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        {
            "email": "centre.coord@sssihl.edu.in",
            "username": "center_coord_user",
            "full_name": "Centre of Excellence Coordinator",
            "role_code": "CENTRE_COORDINATOR",
            "plain_password": "Center_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2018, 6, 1),
            "employee_type": "regular_non_teaching",
        },
        # M8 demo user
        {
            "email": "vc@sssihl.edu.in",
            "username": "vc_user",
            "full_name": "Vice-Chancellor",
            "role_code": "VC",
            "plain_password": "ViceChancellor_Dev1!XZ",
            "gender": "M",
            "joined_on": date(2015, 6, 1),
            "employee_type": "regular_teaching",
        },
    ]
    user_inserted = 0
    for u in users_data:
        role_code = u.pop("role_code")
        plain = u.pop("plain_password")
        is_active = u.pop("is_active", True)
        must_change = u.pop("must_change_password", False)
        # M8 employment fields — default to None/non-teaching for users that don't specify
        gender = u.pop("gender", None)
        joined_on_val = u.pop("joined_on", None)
        employee_type = u.pop("employee_type", "regular_non_teaching")
        new_hash = hash_password(plain)

        stmt = (
            pg_insert(User)
            .values(
                **u,
                password_hash=new_hash,
                is_active=is_active,
                must_change_password=must_change,
                gender=gender,
                joined_on=joined_on_val,
                employee_type=employee_type,
            )
            .on_conflict_do_update(
                constraint="uq_users_email",
                set_={
                    "password_hash": new_hash,
                    "is_active": is_active,
                    "must_change_password": must_change,
                    "failed_login_count": 0,
                    "locked_until": None,
                    "gender": gender,
                    "joined_on": joined_on_val,
                    "employee_type": employee_type,
                },
            )
        )
        result = session.execute(stmt.returning(sa.literal(1).label("x")))
        user_inserted += len(result.fetchall())
        user = session.exec(select(User).where(User.email == u["email"])).one()

        _exec_insert(
            session,
            pg_insert(UserRole)
            .values(user_id=user.id, role_id=roles[role_code].id)
            .on_conflict_do_nothing(),
        )
        # Every user gets BASIC_USER in addition to their primary role.
        if role_code != "BASIC_USER":
            _exec_insert(
                session,
                pg_insert(UserRole)
                .values(user_id=user.id, role_id=roles["BASIC_USER"].id)
                .on_conflict_do_nothing(),
            )
    counts["users"] = user_inserted

    # ── Holidays ──────────────────────────────────────────────────────────────
    holidays_data = [
        {"holiday_date": date(2025, 10, 2),  "name": "Gandhi Jayanti"},
        {"holiday_date": date(2025, 11, 23), "name": "Sai Baba Birthday"},
    ]
    hol_inserted = 0
    for h in holidays_data:
        hol_inserted += _exec_insert(
            session,
            pg_insert(Holiday)
            .values(academic_year_id=ay.id, **h)
            .on_conflict_do_nothing(constraint="uq_holidays_date_ay"),
        )
    # Holidays for 2024-25 (locked AY — demonstrates immutable AY data)
    holidays_prev_data = [
        {"holiday_date": date(2024, 10, 2),  "name": "Gandhi Jayanti"},
        {"holiday_date": date(2024, 11, 23), "name": "Sai Baba Birthday"},
    ]
    for h in holidays_prev_data:
        hol_inserted += _exec_insert(
            session,
            pg_insert(Holiday)
            .values(academic_year_id=ay_prev.id, **h)
            .on_conflict_do_nothing(constraint="uq_holidays_date_ay"),
        )
    counts["holidays"] = hol_inserted

    # ── RoleEmail ─────────────────────────────────────────────────────────────
    # BOOTSTRAP PLACEHOLDERS — sufficient to demonstrate M4 calendar
    # phase-transition emails at the gate.  Real role-email addresses are
    # runtime-managed data configured by Registrar/Reg-office/SysAdmin via
    # the role-email management UI (scheduled ~M5).  Seed is NOT the
    # authoritative address book; these rows exist only so the notification
    # mechanism can be exercised on a fresh DB.
    _role_emails = [
        {"role_code": "SYSTEM_ADMIN",         "email": "admin@example.dev"},
        {"role_code": "IQAC_COORDINATOR",     "email": "iqac@example.dev"},
        {"role_code": "REGISTRAR",            "email": "registrar@example.dev"},
        {"role_code": "DIRECTOR",             "email": "director@example.dev"},
        {"role_code": "DEAN_STUDENT_WELFARE", "email": "dean.sw@example.dev"},
        {"role_code": "HOD",                  "email": "hod.office@example.dev"},
    ]
    re_inserted = 0
    for re_data in _role_emails:
        existing_re = session.exec(
            select(RoleEmail).where(
                RoleEmail.role_code == re_data["role_code"],
                RoleEmail.scope_type.is_(None),  # type: ignore[union-attr]
                RoleEmail.is_deleted == False,  # noqa: E712
            )
        ).first()
        if existing_re:
            if existing_re.email != re_data["email"]:
                existing_re.email = re_data["email"]
                session.add(existing_re)
                re_inserted += 1
        else:
            session.execute(
                pg_insert(RoleEmail).values(
                    role_code=re_data["role_code"],
                    scope_type=None,
                    scope_id=None,
                    email=re_data["email"],
                )
            )
            re_inserted += 1
    counts["role_emails"] = re_inserted

    # ── Placeholder assets (letterhead + template) ────────────────────────────
    import hashlib
    from uuid import uuid4

    from durgam.storage import get_storage_backend

    backend = get_storage_backend()
    sys_admin_user = session.exec(
        select(User).where(User.username == "sys_admin")
    ).first()
    _seed_actor_id = sys_admin_user.id if sys_admin_user else None

    # ── Placeholder DocumentTemplate (letterhead) ──────────────────────────
    # A minimal DOCX so the gate demo has a downloadable letterhead template.
    # DB row creation and file-byte writes are decoupled: uploaded_files/ is
    # gitignored and lost on fresh clone, but the DB persists. The seed must
    # ensure the physical file exists even when the DB row already does.
    from io import BytesIO as _LH_BytesIO

    from docx import Document as _LH_Document

    _lh_doc = _LH_Document()
    _lh_doc.add_paragraph("Placeholder letterhead template — replace via admin UI.")
    _lh_buf = _LH_BytesIO()
    _lh_doc.save(_lh_buf)
    _PLACEHOLDER_DOCX = _lh_buf.getvalue()
    _DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    existing_lh = session.exec(
        select(DocumentTemplate).where(
            DocumentTemplate.purpose == "letterhead",
            DocumentTemplate.role_code == "REGISTRAR",
            DocumentTemplate.is_deleted == False,  # noqa: E712
        )
    ).first()
    lh_inserted = 0
    if existing_lh is None:
        storage_key = uuid4().hex
        sha = hashlib.sha256(_PLACEHOLDER_DOCX).hexdigest()
        backend.put(storage_key, _PLACEHOLDER_DOCX, _DOCX_MIME)

        fa = FileAsset(
            storage_key=storage_key,
            original_name="placeholder_letterhead.docx",
            mime_type=_DOCX_MIME,
            size_bytes=len(_PLACEHOLDER_DOCX),
            sha256=sha,
            owner_user_id=_seed_actor_id,
            purpose="letterhead",
        )
        session.add(fa)
        session.flush()
        session.refresh(fa)

        lh = DocumentTemplate(
            purpose="letterhead",
            role_code="REGISTRAR",
            file_id=fa.id,
            created_by=_seed_actor_id,
            updated_by=_seed_actor_id,
        )
        session.add(lh)
        session.flush()
        lh_inserted = 1
    else:
        fa = session.get(FileAsset, existing_lh.file_id)
        if fa and not backend.exists(fa.storage_key):
            backend.put(fa.storage_key, _PLACEHOLDER_DOCX, _DOCX_MIME)
            log.info("seed_file_restored", key=fa.storage_key, asset="letterhead_REGISTRAR")
    counts["document_templates_letterhead"] = lh_inserted

    # ── Placeholder DocumentTemplate (BoS template) ─────────────────────────
    # A minimal DOCX so the gate demo has a downloadable BoS template.
    # Same decoupled DB-row/file-byte pattern as letterhead above.
    from io import BytesIO as _BytesIO

    from docx import Document as _Document

    _tpl_doc = _Document()
    _tpl_doc.add_heading("Board of Studies — Template", level=1)
    _tpl_doc.add_paragraph("Placeholder template for gate demonstration.")
    _tpl_buf = _BytesIO()
    _tpl_doc.save(_tpl_buf)
    _TPL_DOCX = _tpl_buf.getvalue()
    _TPL_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    existing_tpl = session.exec(
        select(DocumentTemplate).where(
            DocumentTemplate.purpose == "bos",
            DocumentTemplate.role_code.is_(None),  # type: ignore[union-attr]
            DocumentTemplate.is_deleted == False,  # noqa: E712
        )
    ).first()
    tpl_inserted = 0
    if existing_tpl is None:
        tpl_storage_key = uuid4().hex
        tpl_sha = hashlib.sha256(_TPL_DOCX).hexdigest()

        backend.put(tpl_storage_key, _TPL_DOCX, _TPL_MIME)

        tpl_fa = FileAsset(
            storage_key=tpl_storage_key,
            original_name="bos_template.docx",
            mime_type=_TPL_MIME,
            size_bytes=len(_TPL_DOCX),
            sha256=tpl_sha,
            owner_user_id=_seed_actor_id,
            purpose="template",
        )
        session.add(tpl_fa)
        session.flush()
        session.refresh(tpl_fa)

        tpl = DocumentTemplate(
            purpose="bos",
            role_code=None,
            file_id=tpl_fa.id,
            created_by=_seed_actor_id,
            updated_by=_seed_actor_id,
        )
        session.add(tpl)
        session.flush()
        tpl_inserted = 1
    else:
        tpl_fa = session.get(FileAsset, existing_tpl.file_id)
        if tpl_fa and not backend.exists(tpl_fa.storage_key):
            backend.put(tpl_fa.storage_key, _TPL_DOCX, _TPL_MIME)
            log.info("seed_file_restored", key=tpl_fa.storage_key, asset="template_bos")
    counts["document_templates_bos"] = tpl_inserted

    # ── StudentCategoryCount ──────────────────────────────────────────────────
    counts["student_category_counts"] = _exec_insert(
        session,
        pg_insert(StudentCategoryCount)
        .values(
            academic_year_id=ay.id,
            sc_count=120,
            st_count=45,
            obc_count=230,
            ews_count=60,
            general_count=895,
            notes="M0 synthetic seed data",
        )
        .on_conflict_do_nothing(constraint="uq_student_category_counts_ay"),
    )
    # StudentCategoryCount for 2024-25 (locked AY)
    counts["student_category_counts"] += _exec_insert(
        session,
        pg_insert(StudentCategoryCount)
        .values(
            academic_year_id=ay_prev.id,
            sc_count=115,
            st_count=40,
            obc_count=220,
            ews_count=55,
            general_count=870,
            notes="2024-25 historical seed data",
        )
        .on_conflict_do_nothing(constraint="uq_student_category_counts_ay"),
    )

    # ─────────────────────────────────────────────────────────────────────────
    # M3 Organisational Core (§12 M3 gate clause)
    # ─────────────────────────────────────────────────────────────────────────

    # ── Campuses ──────────────────────────────────────────────────────────────
    # Gate: "all four campuses seeded"
    campuses_raw = [
        ("PSN", "Prasanthi Nilayam", "Prasanthi Nilayam, Puttaparthi, Andhra Pradesh 515134"),
        ("BRN", "Brindavan",         "Kadugodi, Whitefield, Bengaluru, Karnataka 560067"),
        ("NDG", "Muddenahalli",      "Muddenahalli, Chikkaballapur, Karnataka 562101"),
        ("ATP", "Anantapur",         "Anantapur, Andhra Pradesh 515001"),
    ]
    campus_inserted = 0
    for code, name, address in campuses_raw:
        campus_inserted += _exec_insert(
            session,
            pg_insert(Campus)
            .values(code=code, name=name, address=address)
            .on_conflict_do_nothing(constraint="uq_campuses_code"),
        )
    counts["campuses"] = campus_inserted
    campuses = {
        c.code: c
        for c in session.exec(
            select(Campus).where(Campus.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── Schools ───────────────────────────────────────────────────────────────
    # Gate: "four schools seeded"; Dean is scoped via UserRole, not per-school column.
    schools_raw = [
        ("SCI", "School of Sciences"),
        ("HSS", "School of Humanities and Social Sciences"),
        ("LL",  "School of Languages and Literature"),
        ("MC",  "School of Management and Commerce"),
    ]
    school_inserted = 0
    for code, name in schools_raw:
        school_inserted += _exec_insert(
            session,
            pg_insert(School)
            .values(code=code, name=name)
            .on_conflict_do_nothing(constraint="uq_schools_code"),
        )
    counts["schools"] = school_inserted
    schools = {
        s.code: s
        for s in session.exec(
            select(School).where(School.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── Departments ───────────────────────────────────────────────────────────
    # Gate: "ten departments seeded"
    # Fields: (code, name, school_code, main_campus_code)
    departments_raw = [
        ("DBIO",  "Biosciences",                    "SCI", "PSN"),
        ("DCHEM", "Chemistry",                      "SCI", "PSN"),
        ("DEDN",  "Education",                      "HSS", "ATP"),
        ("DFNS",  "Food and Nutritional Sciences",  "SCI", "ATP"),
        ("DHSS",  "Humanities and Social Sciences", "HSS", "PSN"),
        ("DLL",   "Languages and Literature",       "LL",  "PSN"),
        ("DMACS", "Mathematics and Computer Science","SCI", "PSN"),
        ("DMC",   "Management and Commerce",        "MC",  "BRN"),
        ("DPHY",  "Physics",                        "SCI", "PSN"),
        ("DPA",   "Performing Arts",                "HSS", "PSN"),
    ]
    dept_inserted = 0
    for code, name, school_code, main_campus_code in departments_raw:
        dept_inserted += _exec_insert(
            session,
            pg_insert(Department)
            .values(
                code=code,
                name=name,
                school_id=schools[school_code].id,
                main_campus_id=campuses[main_campus_code].id,
            )
            .on_conflict_do_nothing(constraint="uq_departments_code"),
        )
    counts["departments"] = dept_inserted
    departments = {
        d.code: d
        for d in session.exec(
            select(Department).where(Department.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── DepartmentCampus ──────────────────────────────────────────────────────
    # (dept_code, campus_code) — per Appendix A campus mapping
    dept_campus_raw = [
        ("DBIO",  "ATP"), ("DBIO",  "PSN"),
        ("DCHEM", "ATP"), ("DCHEM", "PSN"),
        ("DEDN",  "ATP"),
        ("DFNS",  "ATP"),
        ("DHSS",  "ATP"), ("DHSS",  "BRN"), ("DHSS",  "PSN"),
        ("DLL",   "PSN"), ("DLL",   "BRN"), ("DLL",   "NDG"), ("DLL",   "ATP"),
        ("DMACS", "PSN"), ("DMACS", "BRN"), ("DMACS", "NDG"), ("DMACS", "ATP"),
        ("DMC",   "ATP"), ("DMC",   "BRN"),
        ("DPHY",  "ATP"), ("DPHY",  "PSN"),
        ("DPA",   "PSN"),
    ]
    dc_inserted = 0
    for dept_code, campus_code in dept_campus_raw:
        dc_inserted += _exec_insert(
            session,
            pg_insert(DepartmentCampus)
            .values(
                department_id=departments[dept_code].id,
                campus_id=campuses[campus_code].id,
                has_ahod=False,
            )
            .on_conflict_do_nothing(),
        )
    counts["department_campuses"] = dc_inserted

    # ── SubDepartments ────────────────────────────────────────────────────────
    # Gate: "sub-departments seeded" — 9 rows (5 under DHSS, 4 under DLL)
    subdepts_raw = [
        # Under DHSS
        ("SDPHIL", "Philosophy",        "DHSS"),
        ("SDPSY",  "Psychology",        "DHSS"),
        ("SDHIST", "History",           "DHSS"),
        ("SDPOL",  "Political Science", "DHSS"),
        ("SDECON", "Economics",         "DHSS"),
        # Under DLL
        ("SDENG",  "English",           "DLL"),
        ("SDTEL",  "Telugu",            "DLL"),
        ("SDHIN",  "Hindi",             "DLL"),
        ("SDSAN",  "Sanskrit",          "DLL"),
    ]
    subdept_inserted = 0
    for code, name, parent_code in subdepts_raw:
        subdept_inserted += _exec_insert(
            session,
            pg_insert(SubDepartment)
            .values(
                code=code,
                name=name,
                parent_department_id=departments[parent_code].id,
            )
            .on_conflict_do_nothing(constraint="uq_sub_departments_code"),
        )
    counts["sub_departments"] = subdept_inserted
    subdepts = {
        s.code: s
        for s in session.exec(
            select(SubDepartment).where(SubDepartment.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── SubDepartmentCampus ───────────────────────────────────────────────────
    # Per Appendix A: 23 join rows
    subdept_campus_raw = [
        # DHSS sub-departments
        ("SDPHIL", "ATP"),
        ("SDPSY",  "ATP"), ("SDPSY",  "PSN"),
        ("SDHIST", "PSN"), ("SDHIST", "ATP"),
        ("SDPOL",  "PSN"), ("SDPOL",  "ATP"),
        ("SDECON", "BRN"), ("SDECON", "ATP"),
        # DLL sub-departments
        ("SDENG",  "PSN"), ("SDENG",  "BRN"), ("SDENG",  "NDG"), ("SDENG",  "ATP"),
        ("SDTEL",  "PSN"), ("SDTEL",  "ATP"),
        ("SDHIN",  "PSN"), ("SDHIN",  "BRN"), ("SDHIN",  "NDG"), ("SDHIN",  "ATP"),
        ("SDSAN",  "PSN"), ("SDSAN",  "BRN"), ("SDSAN",  "NDG"), ("SDSAN",  "ATP"),
    ]
    sdc_inserted = 0
    for subdept_code, campus_code in subdept_campus_raw:
        sdc_inserted += _exec_insert(
            session,
            pg_insert(SubDepartmentCampus)
            .values(
                sub_department_id=subdepts[subdept_code].id,
                campus_id=campuses[campus_code].id,
            )
            .on_conflict_do_nothing(),
        )
    counts["sub_department_campuses"] = sdc_inserted

    # ── Centres of Excellence ─────────────────────────────────────────────────
    # Gate: "centres seeded" — 4 rows
    centres_raw = [
        ("CMB",  "Centre for Mathematical Biology",          "PSN"),
        ("CSSS", "Centre for Sri Sathya Sai Studies",        "PSN"),
        ("CADS", "Centre for Actuarial and Data Sciences",   "BRN"),
        ("CSD",  "Centre for Sustainable Development",       "PSN"),
    ]
    centre_inserted = 0
    for code, name, campus_code in centres_raw:
        centre_inserted += _exec_insert(
            session,
            pg_insert(CentreOfExcellence)
            .values(code=code, name=name, campus_id=campuses[campus_code].id)
            .on_conflict_do_nothing(constraint="uq_centres_of_excellence_code"),
        )
    counts["centres_of_excellence"] = centre_inserted

    # ── Example Program — BSCMATH ─────────────────────────────────────────────
    # Gate: "one program seeded with full PEO/PO/PSO/regulation/scheme/exit-level data"
    # Program belongs to DMACS; degree type unconstrained string at M3 (OQ-M3-10).
    prog_inserted = _exec_insert(
        session,
        pg_insert(Program)
        .values(
            code="BSCMATH",
            name="Bachelor of Science in Mathematics",
            department_id=departments["DMACS"].id,
            degree_type="BSc",
            duration_years=3,
            is_active=True,
        )
        .on_conflict_do_nothing(constraint="uq_programs_code"),
    )
    counts["programs"] = prog_inserted
    bscmath = session.exec(
        select(Program).where(Program.code == "BSCMATH")
    ).one()

    # ── Seed Courses ──────────────────────────────────────────────────────────
    # Three courses linked to BSCMATH / DMACS.  M13 extends with full taxonomy.
    # Refinement 3 fields: code, name, program_id, department_id, credits,
    # lecture, tutorial, practical, evaluation.
    courses_raw = [
        ("MAT101", "Calculus and Differential Equations", 4, 3, 1, 0, "E"),
        ("MAT102", "Linear Algebra and Abstract Algebra", 4, 3, 1, 0, "E"),
        ("PHY101", "Classical Mechanics",                 4, 3, 0, 2, "IE"),
    ]
    course_inserted = 0
    for code, name, credits, lec, tut, prac, evaluation in courses_raw:
        course_inserted += _exec_insert(
            session,
            pg_insert(Course)
            .values(
                code=code,
                name=name,
                program_id=bscmath.id,
                department_id=departments["DMACS"].id,
                credits=credits,
                lecture=lec,
                tutorial=tut,
                practical=prac,
                evaluation=evaluation,
                is_active=True,
            )
            .on_conflict_do_nothing(constraint="uq_courses_code"),
        )
    counts["courses"] = course_inserted
    courses = {
        c.code: c
        for c in session.exec(
            select(Course).where(Course.is_deleted == False)  # noqa: E712
        ).all()
    }

    # ── ProgramOutcomes: PEOs, POs, PSOs ──────────────────────────────────────
    outcomes_raw = [
        # (outcome_type, code, display_order, description)
        ("PEO", "PEO1", 1,
         "Graduates will apply mathematical foundations to solve complex "
         "scientific and engineering problems."),
        ("PEO", "PEO2", 2,
         "Graduates will contribute to research, industry, or higher education "
         "through sound mathematical reasoning."),
        ("PEO", "PEO3", 3,
         "Graduates will demonstrate professional integrity, teamwork, and "
         "commitment to lifelong learning."),
        ("PO",  "PO1",  1,
         "Apply knowledge of mathematics, statistics, and computing to solve problems."),
        ("PO",  "PO2",  2,
         "Design algorithms and mathematical models for real-world phenomena."),
        ("PO",  "PO3",  3,
         "Analyse and interpret quantitative data using appropriate mathematical tools."),
        ("PO",  "PO4",  4,
         "Formulate and prove mathematical propositions with rigour."),
        ("PO",  "PO5",  5,
         "Communicate mathematical ideas clearly in written and oral form."),
        ("PO",  "PO6",  6,
         "Apply ethical reasoning in professional and academic contexts."),
        ("PSO", "PSO1", 1,
         "Demonstrate proficiency in real analysis, abstract algebra, and topology."),
        ("PSO", "PSO2", 2,
         "Apply numerical methods and computational tools to mathematical problems."),
        ("PSO", "PSO3", 3,
         "Model and solve optimisation problems using linear and nonlinear techniques."),
    ]
    outcome_inserted = 0
    for outcome_type, code, display_order, description in outcomes_raw:
        outcome_inserted += _exec_insert(
            session,
            pg_insert(ProgramOutcome)
            .values(
                program_id=bscmath.id,
                outcome_type=outcome_type,
                code=code,
                description=description,
                display_order=display_order,
            )
            .on_conflict_do_nothing(
                constraint="uq_program_outcomes_program_type_code"
            ),
        )
    counts["program_outcomes"] = outcome_inserted

    # ── ProgramRegulation ─────────────────────────────────────────────────────
    reg_inserted = _exec_insert(
        session,
        pg_insert(ProgramRegulation)
        .values(
            program_id=bscmath.id,
            code="R2021",
            effective_from_year=2021,
            description=(
                "Regulation 2021 — Choice Based Credit System aligned "
                "with NEP 2020 guidelines."
            ),
        )
        .on_conflict_do_nothing(constraint="uq_program_regulations_program_code"),
    )
    counts["program_regulations"] = reg_inserted
    regulation = session.exec(
        select(ProgramRegulation).where(
            ProgramRegulation.program_id == bscmath.id,
            ProgramRegulation.code == "R2021",
        )
    ).one()

    # ── ProgramScheme — Semester 1 ────────────────────────────────────────────
    scheme_inserted = _exec_insert(
        session,
        pg_insert(ProgramScheme)
        .values(
            program_id=bscmath.id,
            regulation_id=regulation.id,
            semester=1,
            total_credits=12,
        )
        .on_conflict_do_nothing(constraint="uq_program_schemes_program_reg_sem"),
    )
    counts["program_schemes"] = scheme_inserted
    scheme = session.exec(
        select(ProgramScheme).where(
            ProgramScheme.program_id == bscmath.id,
            ProgramScheme.regulation_id == regulation.id,
            ProgramScheme.semester == 1,
        )
    ).one()

    # ── ProgramSchemeCourse ───────────────────────────────────────────────────
    psc_inserted = 0
    for course_code in ("MAT101", "MAT102", "PHY101"):
        psc_inserted += _exec_insert(
            session,
            pg_insert(ProgramSchemeCourse)
            .values(scheme_id=scheme.id, course_id=courses[course_code].id)
            .on_conflict_do_nothing(),
        )
    counts["program_scheme_courses"] = psc_inserted

    # ── ProgramSpecialisation ─────────────────────────────────────────────────
    spec_inserted = _exec_insert(
        session,
        pg_insert(ProgramSpecialisation)
        .values(
            program_id=bscmath.id,
            code="APPMATH",
            name="Applied Mathematics",
            description=(
                "Specialisation in numerical analysis, optimisation, "
                "and mathematical modelling."
            ),
        )
        .on_conflict_do_nothing(constraint="uq_program_specialisations_program_code"),
    )
    counts["program_specialisations"] = spec_inserted

    # ── ProgramExitLevel ──────────────────────────────────────────────────────
    exit_levels_raw = [
        ("UG Certificate",    40,  "Exit after completing UG Certificate requirements."),
        ("UG Diploma",        80,  "Exit after completing UG Diploma requirements."),
        ("BSc Mathematics",  120,  "Full three-year BSc programme completion."),
    ]
    el_inserted = 0
    for level_name, required_credits, description in exit_levels_raw:
        el_inserted += _exec_insert(
            session,
            pg_insert(ProgramExitLevel)
            .values(
                program_id=bscmath.id,
                level_name=level_name,
                required_credits=required_credits,
                description=description,
            )
            .on_conflict_do_nothing(
                constraint="uq_program_exit_levels_program_level"
            ),
        )
    counts["program_exit_levels"] = el_inserted

    # ── University Vision/Mission (singleton, E-001) ───────────────────────────
    # Synthetic placeholder text — real text from sponsor at a later milestone.
    # Application-level singleton enforcement (OQ-M3-9): insert if none exists.
    _UNIV_VISION = (
        "To awaken humanity to the divinity within through excellence "
        "in education, inquiry, and character formation."
    )
    uvm = session.exec(
        select(UniversityVisionMission).where(
            UniversityVisionMission.is_deleted == False  # noqa: E712
        )
    ).first()
    if uvm is None:
        session.execute(
            pg_insert(UniversityVisionMission).values(vision=_UNIV_VISION)
        )
        session.flush()
        uvm = session.exec(select(UniversityVisionMission)).first()
        counts["university_vision_missions"] = 1
    else:
        counts["university_vision_missions"] = 0

    assert uvm is not None
    mission_count = session.exec(
        select(sa.func.count(UniversityMission.id)).where(
            UniversityMission.university_vision_id == uvm.id,
            UniversityMission.is_deleted == False,  # noqa: E712
        )
    ).one()
    if mission_count == 0:
        for i, statement in enumerate(
            [
                "To provide value-based education that integrates academic excellence "
                "with spiritual and ethical formation.",
                "To foster a community of scholars committed to the pursuit of truth, "
                "service, and sustainable development.",
                "To serve as a model of educational innovation that bridges ancient "
                "wisdom and modern knowledge.",
            ],
            start=1,
        ):
            session.execute(
                pg_insert(UniversityMission).values(
                    university_vision_id=uvm.id,
                    statement=statement,
                    display_order=i,
                )
            )
        counts["university_missions"] = 3
    else:
        counts["university_missions"] = 0

    # ── DMACS Vision/Mission (E-001) ───────────────────────────────────────────
    _DMACS_VISION = (
        "To nurture mathematical thinkers who bring rigour, creativity, "
        "and purpose to the service of society."
    )
    dvm = session.exec(
        select(DepartmentVisionMission).where(
            DepartmentVisionMission.department_id == departments["DMACS"].id,
            DepartmentVisionMission.is_deleted == False,  # noqa: E712
        )
    ).first()
    if dvm is None:
        session.execute(
            pg_insert(DepartmentVisionMission).values(
                department_id=departments["DMACS"].id,
                vision=_DMACS_VISION,
            )
        )
        session.flush()
        dvm = session.exec(
            select(DepartmentVisionMission).where(
                DepartmentVisionMission.department_id == departments["DMACS"].id
            )
        ).first()
        counts["department_vision_missions"] = 1
    else:
        counts["department_vision_missions"] = 0

    assert dvm is not None
    dept_mission_count = session.exec(
        select(sa.func.count(DepartmentMission.id)).where(
            DepartmentMission.department_vision_id == dvm.id,
            DepartmentMission.is_deleted == False,  # noqa: E712
        )
    ).one()
    if dept_mission_count == 0:
        for i, statement in enumerate(
            [
                "To provide deep training in core mathematical disciplines "
                "and their interdisciplinary applications.",
                "To cultivate a culture of collaborative inquiry, integrity, "
                "and lifelong mathematical learning.",
            ],
            start=1,
        ):
            session.execute(
                pg_insert(DepartmentMission).values(
                    department_vision_id=dvm.id,
                    statement=statement,
                    display_order=i,
                )
            )
        counts["department_missions"] = 2
    else:
        counts["department_missions"] = 0

    # ── ClassTimingsConfig (singleton) ───────────────────────────────────────
    # Default: 8 periods × 50 min, starting 08:00, lunch break after period 4.
    ctc = session.exec(
        select(ClassTimingsConfig).where(ClassTimingsConfig.is_deleted == False)  # noqa: E712
    ).first()
    if ctc is None:
        session.execute(
            pg_insert(ClassTimingsConfig).values(
                periods_per_day=8,
                period_duration_minutes=50,
                first_period_start="08:00",
                break_after_period=4,
                break_duration_minutes=45,
            )
        )
        counts["class_timings_configs"] = 1
    else:
        counts["class_timings_configs"] = 0

    # ── WorkingDaysConfig (singleton) ─────────────────────────────────────────
    wdc = session.exec(
        select(WorkingDaysConfig).where(WorkingDaysConfig.is_deleted == False)  # noqa: E712
    ).first()
    if wdc is None:
        session.execute(pg_insert(WorkingDaysConfig).values(days_per_week=5))
        counts["working_days_configs"] = 1
    else:
        counts["working_days_configs"] = 0

    # ── Scoped role assignments (requires departments to be seeded first) ─────
    # dean_sci → DEAN role scoped to SCI school
    dean_sci_user = session.exec(
        select(User).where(User.username == "dean_sci")
    ).first()
    if dean_sci_user and "SCI" in schools:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=dean_sci_user.id,
                role_id=roles["DEAN"].id,
                scope_type="school",
                scope_id=schools["SCI"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={
                    "scope_type": "school",
                    "scope_id": schools["SCI"].id,
                },
            )
        )

    # hod_dmacs → HOD role scoped to DMACS department
    hod_dmacs_user = session.exec(
        select(User).where(User.username == "hod_dmacs")
    ).first()
    if hod_dmacs_user and "DMACS" in departments:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=hod_dmacs_user.id,
                role_id=roles["HOD"].id,
                scope_type="department",
                scope_id=departments["DMACS"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={
                    "scope_type": "department",
                    "scope_id": departments["DMACS"].id,
                },
            )
        )

    # director_psn → DIRECTOR role scoped to PSN campus (M4)
    director_psn_user = session.exec(
        select(User).where(User.username == "director_psn")
    ).first()
    if director_psn_user and "PSN" in campuses:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=director_psn_user.id,
                role_id=roles["DIRECTOR"].id,
                scope_type="campus",
                scope_id=campuses["PSN"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={
                    "scope_type": "campus",
                    "scope_id": campuses["PSN"].id,
                },
            )
        )

    # hod_office_dmacs → HOD_OFFICE role scoped to DMACS department (M5b-R2)
    _hod_office_user = session.exec(
        select(User).where(User.username == "hod_office_dmacs")
    ).first()
    if _hod_office_user and "DMACS" in departments:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=_hod_office_user.id,
                role_id=roles["HOD_OFFICE"].id,
                scope_type="department",
                scope_id=departments["DMACS"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={"scope_type": "department", "scope_id": departments["DMACS"].id},
            )
        )

    # ahod_dmacs → AHOD role scoped to DMACS department (M5b-R2)
    _ahod_user = session.exec(
        select(User).where(User.username == "ahod_dmacs")
    ).first()
    if _ahod_user and "DMACS" in departments:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=_ahod_user.id,
                role_id=roles["AHOD"].id,
                scope_type="department",
                scope_id=departments["DMACS"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={"scope_type": "department", "scope_id": departments["DMACS"].id},
            )
        )

    # director_office_psn → DIRECTOR_OFFICE role scoped to PSN campus (M5b-R2)
    _dir_office_user = session.exec(
        select(User).where(User.username == "director_office_psn")
    ).first()
    if _dir_office_user and "PSN" in campuses:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=_dir_office_user.id,
                role_id=roles["DIRECTOR_OFFICE"].id,
                scope_type="campus",
                scope_id=campuses["PSN"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={"scope_type": "campus", "scope_id": campuses["PSN"].id},
            )
        )

    # deputy_director_psn → DEPUTY_DIRECTOR role scoped to PSN campus (M5b-R2)
    _dep_dir_user = session.exec(
        select(User).where(User.username == "deputy_director_psn")
    ).first()
    if _dep_dir_user and "PSN" in campuses:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=_dep_dir_user.id,
                role_id=roles["DEPUTY_DIRECTOR"].id,
                scope_type="campus",
                scope_id=campuses["PSN"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={"scope_type": "campus", "scope_id": campuses["PSN"].id},
            )
        )

    # faculty_user → FACULTY role scoped to DMACS department (M5b-R2)
    _faculty_user = session.exec(
        select(User).where(User.username == "faculty_user")
    ).first()
    if _faculty_user and "DMACS" in departments:
        session.execute(
            pg_insert(UserRole)
            .values(
                user_id=_faculty_user.id,
                role_id=roles["FACULTY"].id,
                scope_type="department",
                scope_id=departments["DMACS"].id,
            )
            .on_conflict_do_update(
                index_elements=["user_id", "role_id"],
                set_={"scope_type": "department", "scope_id": departments["DMACS"].id},
            )
        )

    # ─────────────────────────────────────────────────────────────────────────
    # M4 Sample Calendar Entries (for 2025-26, unlocked AY)
    # ─────────────────────────────────────────────────────────────────────────
    registrar_user = session.exec(
        select(User).where(User.username == "registrar_user")
    ).first()
    existing_cal_count = session.exec(
        select(sa.func.count(CalendarEntry.id)).where(
            CalendarEntry.academic_year_id == ay.id,
            CalendarEntry.is_deleted == False,  # noqa: E712
        )
    ).one()
    if registrar_user and existing_cal_count == 0:
        cal_entries_raw = [
            {
                "title": "Semester 1 Begins",
                "entry_type": "sem_begin",
                "starts_at": datetime(2025, 7, 14, 9, 0, tzinfo=UTC),
                "ends_at": datetime(2025, 7, 14, 17, 0, tzinfo=UTC),
                "owner_user_id": registrar_user.id,
                "owner_role_code": "REGISTRAR",
            },
            {
                "title": "Semester 1 Ends",
                "entry_type": "sem_end",
                "starts_at": datetime(2025, 11, 28, 9, 0, tzinfo=UTC),
                "ends_at": datetime(2025, 11, 28, 17, 0, tzinfo=UTC),
                "owner_user_id": registrar_user.id,
                "owner_role_code": "REGISTRAR",
            },
            {
                "title": "CIE-1",
                "entry_type": "cie",
                "starts_at": datetime(2025, 9, 8, 9, 0, tzinfo=UTC),
                "ends_at": datetime(2025, 9, 12, 17, 0, tzinfo=UTC),
                "owner_user_id": registrar_user.id,
                "owner_role_code": "REGISTRAR",
            },
        ]
        for entry in cal_entries_raw:
            session.execute(pg_insert(CalendarEntry).values(academic_year_id=ay.id, **entry))
        counts["calendar_entries"] = 3
    else:
        counts["calendar_entries"] = 0

    # ─────────────────────────────────────────────────────────────────────────
    # M5b Session 7: Purchase Policy & Approval Config (E-007)
    # ─────────────────────────────────────────────────────────────────────────

    # ── Designation vocabulary (§8.3 — M10 Phase 1B expanded taxonomy) ─────────
    # Legacy codes (senior_professor, professor, associate_professor,
    # assistant_professor) are soft-deleted by migration cb2de963f0b8.
    # This seed inserts only the 7 new codes (idempotent via ON CONFLICT DO NOTHING).
    designations_raw = [
        ("sr_prof",       "Senior Professor",                        1),
        ("prof",          "Professor",                               2),
        ("assoc_prof",    "Associate Professor",                     3),
        ("asst_prof_l10", "Assistant Professor (Academic Level 10)", 4),
        ("asst_prof_l11", "Assistant Professor (Academic Level 11)", 5),
        ("asst_prof_l12", "Assistant Professor (Academic Level 12)", 6),
        ("instructor",    "Instructor",                              7),
    ]
    desig_inserted = 0
    for code, name, rank in designations_raw:
        desig_inserted += _exec_insert(
            session,
            pg_insert(Designation)
            .values(code=code, name=name, rank=rank)
            .on_conflict_do_nothing(constraint="uq_designations_code"),
        )
    counts["designations"] = desig_inserted

    # ── PurchaseProcedureRule (10 rows: 5 tiers × 2 fund sources) ────────────
    # Source: docs/4384-Purchase Procedures of the Institute.PDF
    # Institute Budgeted Funds tiers:
    ppr_institute = [
        {
            "fund_source": "institute", "tier": 1,
            "floor_amount": 0, "ceiling_amount": 10_000,
            "min_quotes_required": False, "min_quote_count": 0,
            "quote_at_discretion": True,
            "comparative_statement_required": False,
            "approving_authority_role_codes": ["DIRECTOR"],
            "committee_level": None,
        },
        {
            "fund_source": "institute", "tier": 2,
            "floor_amount": 10_001, "ceiling_amount": 50_000,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["DIRECTOR"],
            "committee_level": None,
        },
        {
            "fund_source": "institute", "tier": 3,
            "floor_amount": 50_001, "ceiling_amount": 499_999,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["REGISTRAR"],
            "committee_level": "campus_purchase_committee",
        },
        {
            "fund_source": "institute", "tier": 4,
            "floor_amount": 500_000, "ceiling_amount": 999_999,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["VC"],
            "committee_level": "central_purchase_committee",
        },
        {
            "fund_source": "institute", "tier": 5,
            "floor_amount": 1_000_000, "ceiling_amount": None,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            # BOM (Board of Management) is an offline statutory body;
            # M7 handles this as offline approval. Stored as literal string,
            # no FK/join to the roles table.
            "approving_authority_role_codes": ["BOM"],
            "committee_level": "central_purchase_committee",
        },
    ]
    # Projects / UGC Funds tiers:
    ppr_projects = [
        {
            "fund_source": "projects_ugc", "tier": 1,
            "floor_amount": 0, "ceiling_amount": 10_000,
            "min_quotes_required": False, "min_quote_count": 0,
            "quote_at_discretion": True,
            "comparative_statement_required": False,
            "approving_authority_role_codes": ["HOD", "DIRECTOR", "DEAN"],
            "committee_level": None,
        },
        {
            # PDF states tier-2 floor as 10,000 (overlapping with tier-1 ceiling).
            # Normalized to 10,001 to avoid range overlap.
            "fund_source": "projects_ugc", "tier": 2,
            "floor_amount": 10_001, "ceiling_amount": 50_000,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["HOD", "DIRECTOR", "DEAN"],
            "committee_level": None,
        },
        {
            # PDF states tier-3 floor as 50,000; normalized to 50,001.
            "fund_source": "projects_ugc", "tier": 3,
            "floor_amount": 50_001, "ceiling_amount": 499_999,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["REGISTRAR"],
            "committee_level": "campus_purchase_committee",
        },
        {
            "fund_source": "projects_ugc", "tier": 4,
            "floor_amount": 500_000, "ceiling_amount": 999_999,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["VC"],
            "committee_level": "central_purchase_committee",
        },
        {
            "fund_source": "projects_ugc", "tier": 5,
            "floor_amount": 1_000_000, "ceiling_amount": None,
            "min_quotes_required": True, "min_quote_count": 3,
            "quote_at_discretion": False,
            "comparative_statement_required": True,
            "approving_authority_role_codes": ["VC"],
            "committee_level": "central_purchase_committee",
        },
    ]
    ppr_inserted = 0
    for rule in ppr_institute + ppr_projects:
        ppr_inserted += _exec_insert(
            session,
            pg_insert(PurchaseProcedureRule)
            .values(**rule)
            .on_conflict_do_nothing(constraint="uq_ppr_fund_source_tier"),
        )
    counts["purchase_procedure_rules"] = ppr_inserted

    # ── PurchaseCommitteeTemplate (2 rows) ───────────────────────────────────
    # Campus purchase committee (E-007):
    # - Director is NOT a member (director_excluded=True) — Director gives
    #   comments and forwards. Per PDF: "Director should not be a member of the
    #   committee as he/she would be giving comments/recommendation and
    #   forwarding the proposal to Ad-block for approval."
    #   Forwarding topology: Director → Registrar → VC (captured in notes;
    #   M7 runtime enforces the actual routing).
    # - Three faculty members from different departments + HoD of concerned dept.
    # - eligible_designations is rank-ordered (highest rank first).
    pct_campus_inserted = _exec_insert(
        session,
        pg_insert(PurchaseCommitteeTemplate)
        .values(
            committee_type="campus_purchase_committee",
            eligible_designations=[
                "sr_prof", "prof", "assoc_prof",
            ],
            faculty_member_count=3,
            members_from_different_departments=True,
            fixed_role_members=["HOD"],
            director_excluded=True,
            escalation_designate_role_code=None,
            external_expert_mode="proxied_with_proof",
            topology="concurrent",
            notes=(
                "Director is NOT a member. Director gives comments/recommendation "
                "and forwards to Registrar → VC. Three faculty members from "
                "departments other than the initiating department."
            ),
        )
        .on_conflict_do_nothing(constraint="uq_pct_committee_type"),
    )
    # Central purchase committee (E-007):
    # - Registrar is escalation designate (takes proposal to VC).
    # - Three faculty from different depts + HoD + Finance Officer + Registrar.
    pct_central_inserted = _exec_insert(
        session,
        pg_insert(PurchaseCommitteeTemplate)
        .values(
            committee_type="central_purchase_committee",
            eligible_designations=[
                "sr_prof", "prof", "assoc_prof",
            ],
            faculty_member_count=3,
            members_from_different_departments=True,
            fixed_role_members=["HOD", "FINANCE_OFFICER", "REGISTRAR"],
            director_excluded=False,
            escalation_designate_role_code="REGISTRAR",
            external_expert_mode="proxied_with_proof",
            topology="concurrent",
            notes=(
                "Registrar takes proposal to VC. Three faculty members from "
                "departments other than the initiating department, plus HoD of "
                "concerned dept, Finance Officer, and Registrar."
            ),
        )
        .on_conflict_do_nothing(constraint="uq_pct_committee_type"),
    )
    counts["purchase_committee_templates"] = pct_campus_inserted + pct_central_inserted

    # ── ApprovalProcess — CPC_FUND_RELEASE ───────────────────────────────────
    cpc_inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code="CPC_FUND_RELEASE",
            title="Central Purchase Committee Fund Release",
            requestor_role_codes=["HOD", "AHOD"],
            channel_role_codes=[
                "REGISTRAR", "FINANCE_OFFICER", "CPC_CHAIRPERSON", "VC",
            ],
            is_finance=True,
        )
        .on_conflict_do_nothing(constraint="uq_approval_processes_code"),
    )
    counts["approval_processes"] = cpc_inserted

    # ── ApprovalProcess — NRF_APPROVAL ──────────────────────────────────────
    nrf_inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code="NRF_APPROVAL",
            title="Non-Regular Faculty Approval",
            requestor_role_codes=["HOD", "AHOD"],
            channel_role_codes=["HOD", "DEAN", "REGISTRAR"],
            requires_upward_attachments=True,
            max_upward_attachments=5,
            requires_downward_attachments=False,
            max_downward_attachments=3,
            is_finance=False,
        )
        .on_conflict_do_update(
            constraint="uq_approval_processes_code",
            set_={"channel_role_codes": ["HOD", "DEAN", "REGISTRAR"]},
        ),
    )
    counts["approval_processes"] += nrf_inserted

    # ── ApprovalProcess — DSW_CLEARANCE ─────────────────────────────────────
    # Demo process so dean_sw (DEAN_STUDENT_WELFARE role) appears as a channel
    # approver and sees the Approvals nav link via dynamic_check.
    dsw_inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code="DSW_CLEARANCE",
            title="Dean Student Welfare Clearance",
            requestor_role_codes=["HOD", "AHOD"],
            channel_role_codes=["HOD", "DEAN_STUDENT_WELFARE", "REGISTRAR"],
            requires_upward_attachments=False,
            max_upward_attachments=3,
            requires_downward_attachments=False,
            max_downward_attachments=0,
            is_finance=False,
        )
        .on_conflict_do_nothing(constraint="uq_approval_processes_code"),
    )
    counts["approval_processes"] += dsw_inserted

    # ── ApprovalProcess — LEAVE_APPROVAL (M8) ───────────────────────────────
    # Generic leave-request approval. Per-request channel is resolved at
    # submit-time via LeaveSanctionAuthorityRule (Path A architecture).
    # channel_role_codes = union of all sanctioner roles in the leave matrix;
    # this drives is_channel_approver() nav gating without hard-coding any
    # specific channel — the actual approval channel is in resolved_channel_json
    # on each individual ApprovalRequest row.
    leave_inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code="LEAVE_APPROVAL",
            title="Leave Approval",
            # Any authenticated user may apply for leave; requestor gating is via
            # the leave UI, not the ApprovalProcess template.
            requestor_role_codes=None,
            # Union of all sanctioner roles in the leave matrix — drives
            # is_channel_approver() nav gating so these roles see the Approvals link.
            channel_role_codes=[
                "DIRECTOR", "VC", "REGISTRAR",
                "FINANCE_OFFICER", "CONTROLLER_OF_EXAMINATIONS",
            ],
            informational_cc_role_codes=["HR_HEAD"],
            is_finance=False,
        )
        .on_conflict_do_update(
            constraint="uq_approval_processes_code",
            set_={
                "channel_role_codes": [
                    "DIRECTOR", "VC", "REGISTRAR",
                    "FINANCE_OFFICER", "CONTROLLER_OF_EXAMINATIONS",
                ],
                "informational_cc_role_codes": ["HR_HEAD"],
            },
        ),
    )
    counts["approval_processes"] += leave_inserted

    # ── LeaveSanctionAuthorityRule — load from YAML (M8) ─────────────────────
    # Idempotent: upsert on natural key; soft-deletes orphans.
    from pathlib import Path as _Path

    from durgam.repositories.leave import LeaveSanctionRuleRepository
    from durgam.services.leave_sanction_rule import LeaveSanctionRuleService

    _leave_repo = LeaveSanctionRuleRepository(session)
    _leave_svc = LeaveSanctionRuleService(session, _leave_repo)
    _yaml_path = _Path(__file__).parent.parent / "seeds" / "leave_sanction_matrix.yaml"
    _leave_actor = sys_admin_user.id if sys_admin_user else _seed_actor_id
    _matrix_counts = _leave_svc.load_from_yaml(_yaml_path, actor_id=_leave_actor)
    counts["leave_sanction_rules_inserted"] = _matrix_counts["inserted"]
    counts["leave_sanction_rules_updated"] = _matrix_counts["updated"]
    counts["leave_sanction_rules_orphaned"] = _matrix_counts["orphaned_soft_deleted"]
    log.info(
        "leave_matrix_seeded",
        inserted=_matrix_counts["inserted"],
        updated=_matrix_counts["updated"],
        orphaned=_matrix_counts["orphaned_soft_deleted"],
    )

    # ── AnnouncementCategory (M9) ──────────────────────────────────────────────
    # Default 9 categories per M9.md Q17 freeze. Operational additions via
    # Registrar-tier admin UI post-launch.
    categories_data = [
        {"code": "CIRCULAR",     "name": "Circular",          "display_order": 10, "is_active": True},
        {"code": "ORDER",        "name": "Order",             "display_order": 20, "is_active": True},
        {"code": "NOTICE",       "name": "Notice",            "display_order": 30, "is_active": True},
        {"code": "NOTIFICATION", "name": "Notification",      "display_order": 40, "is_active": True},
        {"code": "MEMORANDUM",   "name": "Office Memorandum", "display_order": 50, "is_active": True},
        {"code": "INVITATION",   "name": "Invitation",        "display_order": 60, "is_active": True},
        {"code": "RESULT",       "name": "Result",            "display_order": 70, "is_active": True},
        {"code": "ADVISORY",     "name": "Advisory",          "display_order": 80, "is_active": True},
        {"code": "GENERAL",      "name": "General",           "display_order": 90, "is_active": True},
    ]
    cat_inserted = 0
    for c in categories_data:
        cat_inserted += _exec_insert(
            session,
            pg_insert(AnnouncementCategory)
            .values(**c)
            .on_conflict_do_nothing(constraint="uq_announcement_categories_code"),
        )
    counts["announcement_categories"] = cat_inserted

    # ── AnnouncementComposerConfig (M9) ────────────────────────────────────────
    # 19 rows per M9.md Q11 freeze. All role_code values verified in seed.py roles list.
    composer_config_data = [
        {"role_code": "VC",                           "priority_rank": 10,  "scope_restriction": None,         "enabled": True},
        {"role_code": "VC_OFFICE",                    "priority_rank": 20,  "scope_restriction": None,         "enabled": True},
        {"role_code": "REGISTRAR",                    "priority_rank": 30,  "scope_restriction": None,         "enabled": True},
        {"role_code": "REGISTRAR_OFFICE",             "priority_rank": 40,  "scope_restriction": None,         "enabled": True},
        {"role_code": "HR_HEAD",                      "priority_rank": 50,  "scope_restriction": None,         "enabled": True},
        {"role_code": "IQAC_COORDINATOR",             "priority_rank": 60,  "scope_restriction": None,         "enabled": True},
        {"role_code": "DEAN",                         "priority_rank": 70,  "scope_restriction": "school",     "enabled": True},
        {"role_code": "DEAN_STUDENT_WELFARE",         "priority_rank": 70,  "scope_restriction": None,         "enabled": True},
        {"role_code": "DEAN_ACADEMIC_AFFAIRS",        "priority_rank": 70,  "scope_restriction": None,         "enabled": True},
        {"role_code": "DEAN_STUDENT_WELFARE_OFFICE",  "priority_rank": 71,  "scope_restriction": None,         "enabled": True},
        {"role_code": "DEAN_ACADEMIC_AFFAIRS_OFFICE", "priority_rank": 71,  "scope_restriction": None,         "enabled": True},
        {"role_code": "DIRECTOR",                     "priority_rank": 80,  "scope_restriction": "campus",     "enabled": True},
        {"role_code": "DIRECTOR_OFFICE",              "priority_rank": 90,  "scope_restriction": "campus",     "enabled": True},
        {"role_code": "CONTROLLER_OF_EXAMINATIONS",   "priority_rank": 100, "scope_restriction": None,         "enabled": True},
        {"role_code": "FINANCE_OFFICER",              "priority_rank": 110, "scope_restriction": None,         "enabled": True},
        {"role_code": "PLACEMENT_OFFICER",            "priority_rank": 120, "scope_restriction": None,         "enabled": True},
        {"role_code": "HOD",                          "priority_rank": 130, "scope_restriction": "department", "enabled": True},
        {"role_code": "CESRC_COORDINATOR",            "priority_rank": 140, "scope_restriction": None,         "enabled": True},
        {"role_code": "CENTRE_COORDINATOR",           "priority_rank": 150, "scope_restriction": "centre",     "enabled": True},
    ]
    ccfg_inserted = 0
    for cc in composer_config_data:
        ccfg_inserted += _exec_insert(
            session,
            pg_insert(AnnouncementComposerConfig)
            .values(**cc)
            .on_conflict_do_nothing(constraint="uq_announcement_composer_configs_role_code"),
        )
    counts["announcement_composer_configs"] = ccfg_inserted

    # ── AudienceGroup (M9) ─────────────────────────────────────────────────────
    # 23 explicit rows per M9.md Q18 freeze + dynamic per-campus EVERYONE_<code> rows.
    # All school/centre codes verified to exist in DB (SCI/HSS/LL/MC; CMB/CSSS/CADS/CSD).
    # STUDENT_UG/PG/PHD program_degree_types subject to TD-043 (resolves empty until
    # User-Program enrollment link is built).
    audience_groups_explicit = [
        # All-hands
        {"code": "ALL", "name": "Everyone", "description": "All active users.",
         "filter_json": {}, "is_active": True},
        # Faculty
        {"code": "FACULTY_ALL", "name": "All Faculty",
         "filter_json": {"role_codes": ["FACULTY", "PROFESSOR", "ASSOC_PROFESSOR"]},
         "is_active": True, "description": None},
        {"code": "FACULTY_SCI", "name": "Sciences Faculty",
         "filter_json": {"role_codes": ["FACULTY", "PROFESSOR", "ASSOC_PROFESSOR"], "scope_type": "school", "scope_codes": ["SCI"]},
         "is_active": True, "description": None},
        {"code": "FACULTY_HSS", "name": "Humanities & Social Sciences Faculty",
         "filter_json": {"role_codes": ["FACULTY", "PROFESSOR", "ASSOC_PROFESSOR"], "scope_type": "school", "scope_codes": ["HSS"]},
         "is_active": True, "description": None},
        {"code": "FACULTY_LL", "name": "Languages & Literature Faculty",
         "filter_json": {"role_codes": ["FACULTY", "PROFESSOR", "ASSOC_PROFESSOR"], "scope_type": "school", "scope_codes": ["LL"]},
         "is_active": True, "description": None},
        {"code": "FACULTY_MC", "name": "Management & Commerce Faculty",
         "filter_json": {"role_codes": ["FACULTY", "PROFESSOR", "ASSOC_PROFESSOR"], "scope_type": "school", "scope_codes": ["MC"]},
         "is_active": True, "description": None},
        {"code": "PROFESSORS_ALL", "name": "All Professors",
         "filter_json": {"role_codes": ["PROFESSOR"]},
         "is_active": True, "description": None},
        # Students — STUDENT_UG/PG/PHD subject to TD-043
        {"code": "STUDENT_ALL", "name": "All Students",
         "filter_json": {"role_codes": ["STUDENT"]},
         "is_active": True, "description": None},
        {"code": "STUDENT_UG", "name": "Undergraduate Students",
         "filter_json": {"role_codes": ["STUDENT"], "program_degree_types": ["BSc", "BA", "BCom", "BTech", "BBA", "BPharm", "MBBS"]},
         "is_active": True, "description": "Subject to TD-043: resolves empty until User-Program link exists."},
        {"code": "STUDENT_PG", "name": "Postgraduate Students",
         "filter_json": {"role_codes": ["STUDENT"], "program_degree_types": ["MSc", "MA", "MCom", "MBA", "MTech", "MPharm", "MD", "MS"]},
         "is_active": True, "description": "Subject to TD-043: resolves empty until User-Program link exists."},
        {"code": "STUDENT_PHD", "name": "PhD Scholars",
         "filter_json": {"role_codes": ["STUDENT"], "program_degree_types": ["PhD", "DPhil"]},
         "is_active": True, "description": "Subject to TD-043: resolves empty until User-Program link exists."},
        {"code": "STUDENT_SCI", "name": "Sciences Students",
         "filter_json": {"role_codes": ["STUDENT"], "scope_type": "school", "scope_codes": ["SCI"]},
         "is_active": True, "description": None},
        {"code": "STUDENT_HSS", "name": "Humanities & Social Sciences Students",
         "filter_json": {"role_codes": ["STUDENT"], "scope_type": "school", "scope_codes": ["HSS"]},
         "is_active": True, "description": None},
        {"code": "STUDENT_LL", "name": "Languages & Literature Students",
         "filter_json": {"role_codes": ["STUDENT"], "scope_type": "school", "scope_codes": ["LL"]},
         "is_active": True, "description": None},
        {"code": "STUDENT_MC", "name": "Management & Commerce Students",
         "filter_json": {"role_codes": ["STUDENT"], "scope_type": "school", "scope_codes": ["MC"]},
         "is_active": True, "description": None},
        # Centres of Excellence
        {"code": "CENTRE_CMB",  "name": "Centre of Excellence — CMB",
         "filter_json": {"scope_type": "centre", "scope_codes": ["CMB"]},
         "is_active": True, "description": None},
        {"code": "CENTRE_CSSS", "name": "Centre of Excellence — CSSS",
         "filter_json": {"scope_type": "centre", "scope_codes": ["CSSS"]},
         "is_active": True, "description": None},
        {"code": "CENTRE_CADS", "name": "Centre of Excellence — CADS",
         "filter_json": {"scope_type": "centre", "scope_codes": ["CADS"]},
         "is_active": True, "description": None},
        {"code": "CENTRE_CSD",  "name": "Centre of Excellence — CSD",
         "filter_json": {"scope_type": "centre", "scope_codes": ["CSD"]},
         "is_active": True, "description": None},
        # Leadership / staff
        {"code": "HODS_ALL", "name": "All HoDs",
         "filter_json": {"role_codes": ["HOD"]},
         "is_active": True, "description": None},
        {"code": "DEANS_ALL", "name": "All Deans (incl. office variants)",
         "filter_json": {"role_codes": ["DEAN", "DEAN_STUDENT_WELFARE", "DEAN_ACADEMIC_AFFAIRS", "DEAN_STUDENT_WELFARE_OFFICE", "DEAN_ACADEMIC_AFFAIRS_OFFICE"]},
         "is_active": True, "description": None},
        {"code": "DIRECTORS_ALL", "name": "All Directors (incl. office)",
         "filter_json": {"role_codes": ["DIRECTOR", "DIRECTOR_OFFICE"]},
         "is_active": True, "description": None},
        {"code": "OFFICE_STAFF_ALL", "name": "All Office Staff",
         "filter_json": {"role_codes": ["VC_OFFICE", "REGISTRAR_OFFICE", "DIRECTOR_OFFICE", "DEAN_STUDENT_WELFARE_OFFICE", "DEAN_ACADEMIC_AFFAIRS_OFFICE", "HR_OFFICE"]},
         "is_active": True, "description": None},
    ]
    ag_inserted = 0
    for ag in audience_groups_explicit:
        ag_inserted += _exec_insert(
            session,
            pg_insert(AudienceGroup)
            .values(**ag)
            .on_conflict_do_nothing(constraint="uq_audience_groups_code"),
        )

    # Dynamic per-campus EVERYONE_<code> rows — read campuses seeded earlier this run
    all_campuses = session.exec(
        select(Campus).where(Campus.is_deleted == False)  # noqa: E712
    ).all()
    for campus in all_campuses:
        row = {
            "code": f"EVERYONE_{campus.code}",
            "name": f"Everyone in {campus.name}",
            "description": f"All users assigned to the {campus.name} campus.",
            "filter_json": {"scope_type": "campus", "scope_codes": [campus.code]},
            "is_active": True,
        }
        ag_inserted += _exec_insert(
            session,
            pg_insert(AudienceGroup)
            .values(**row)
            .on_conflict_do_nothing(constraint="uq_audience_groups_code"),
        )
    counts["audience_groups"] = ag_inserted

    # ── LeaveCreditPolicy — CL row (M8.1 TD-036) ─────────────────────────────
    # Bootstrap placeholder — real entitlement values managed via admin UI.
    # vacation_entitlement=10 (teaching) / non_vacation_entitlement=12 (non-teaching)
    # per §XXVIII clause 14.
    cl_policy_stmt = (
        pg_insert(LeaveCreditPolicy)
        .values(
            leave_type="CL",
            vacation_entitlement=10.0,
            non_vacation_entitlement=12.0,
            enabled=True,
        )
        .on_conflict_do_nothing(constraint="uq_leave_credit_policies_leave_type")
    )
    counts["leave_credit_policies"] = _exec_insert(session, cl_policy_stmt)

    # ── Faculty seed backfill (M10 Phase 1B) ─────────────────────────────────
    counts["faculty_backfill"] = _seed_faculty_backfill(session)

    # ── ApprovalProcess — faculty_noc (M10 Phase 5B) ─────────────────────────
    counts["faculty_noc_process"] = _seed_faculty_noc_process(session)

    # ── ApprovalProcess — 4 linear faculty processes (M10 Phase 5D) ──────────
    counts["faculty_invited_talk_process"] = _seed_faculty_simple_linear_process(
        session,
        code="faculty_invited_talk",
        title="Faculty Invited Talk Request",
        channel_role_codes=["HOD", "DIRECTOR"],
    )
    counts["faculty_professional_membership_process"] = _seed_faculty_simple_linear_process(
        session,
        code="faculty_professional_membership",
        title="Faculty Professional Body Membership Request",
        channel_role_codes=["HOD", "DIRECTOR"],
    )
    counts["faculty_wfh_process"] = _seed_faculty_simple_linear_process(
        session,
        code="faculty_wfh",
        title="Faculty Work From Home Request",
        channel_role_codes=["HOD", "DIRECTOR"],
    )
    counts["faculty_field_visit_process"] = _seed_faculty_simple_linear_process(
        session,
        code="faculty_field_visit",
        title="Faculty Field/Industry Visit Request",
        channel_role_codes=["HOD", "DIRECTOR"],
    )

    session.commit()
    return counts


def _seed_faculty_backfill(session: Session) -> int:
    """Create Faculty rows for the 7 regular_teaching seeded users.

    Idempotent: upsert on uq_faculties_employee_id (ON CONFLICT DO NOTHING).
    If any referenced user/dept/campus/designation is absent, raises KeyError
    to surface the gap rather than silently skipping rows.

    Mapping frozen in M10 Phase 1B prompt (2026-06-14, Bala authority).
    """
    from datetime import date as _date

    # Resolve FK UUIDs from already-seeded rows.
    dmacs = session.exec(
        select(Department).where(Department.code == "DMACS", Department.is_deleted == False)  # noqa: E712
    ).first()
    if dmacs is None:
        raise KeyError("Dept DMACS not found in seed — cannot backfill Faculty rows")

    psn = session.exec(
        select(Campus).where(Campus.code == "PSN", Campus.is_deleted == False)  # noqa: E712
    ).first()
    if psn is None:
        raise KeyError("Campus PSN not found in seed — cannot backfill Faculty rows")

    # Designation lookup: code -> Designation row (active only)
    desig_rows = session.exec(
        select(Designation).where(Designation.is_deleted == False)  # noqa: E712
    ).all()
    desig_map = {d.code: d for d in desig_rows}

    # Backfill spec: (username, employee_id, desig_code, first_name, last_name,
    #                  phone, ec_name, ec_relation, ec_phone, is_phd, joining_date)
    backfill = [
        ("vc_user",             "DEV-FAC-0001", "sr_prof",      "Vc",      "DevUser",  "9000000001", "Test Contact 1", "Spouse", "9000000101", True,  _date(2010, 6, 1)),
        ("dean_sci",            "DEV-FAC-0002", "prof",         "Dean",    "DevSci",   "9000000002", "Test Contact 2", "Spouse", "9000000102", True,  _date(2012, 6, 1)),
        ("director_psn",        "DEV-FAC-0003", "prof",         "Director","DevPsn",   "9000000003", "Test Contact 3", "Spouse", "9000000103", True,  _date(2013, 6, 1)),
        ("hod_dmacs",           "DEV-FAC-0004", "prof",         "Hod",     "DevDmacs", "9000000004", "Test Contact 4", "Spouse", "9000000104", True,  _date(2014, 6, 1)),
        ("ahod_dmacs",          "DEV-FAC-0005", "assoc_prof",   "Ahod",    "DevDmacs", "9000000005", "Test Contact 5", "Spouse", "9000000105", True,  _date(2016, 6, 1)),
        ("deputy_director_psn", "DEV-FAC-0006", "assoc_prof",   "Deputy",  "DevPsn",   "9000000006", "Test Contact 6", "Spouse", "9000000106", True,  _date(2017, 6, 1)),
        ("faculty_user",        "DEV-FAC-0007", "asst_prof_l10","Faculty", "DevUser",  "9000000007", "Test Contact 7", "Spouse", "9000000107", False, _date(2022, 6, 1)),
    ]

    inserted = 0
    for username, emp_id, desig_code, first_name, last_name, phone, ec_name, ec_rel, ec_phone, is_phd, joining_date in backfill:
        user = session.exec(
            select(User).where(User.username == username, User.is_deleted == False)  # noqa: E712
        ).first()
        if user is None:
            raise KeyError(f"User {username!r} not found — cannot backfill Faculty row")

        desig = desig_map.get(desig_code)
        if desig is None:
            raise KeyError(f"Designation {desig_code!r} not found — cannot backfill Faculty row for {username!r}")

        title = "Dr." if is_phd else "Mr."
        inserted += _exec_insert(
            session,
            pg_insert(Faculty).values(
                user_id=user.id,
                employee_id=emp_id,
                title=title,
                first_name=first_name,
                last_name=last_name,
                designation_id=desig.id,
                department_id=dmacs.id,
                campus_id=psn.id,
                joining_date=joining_date,
                is_vacation_employee=True,
                phone=phone,
                emergency_contact_name=ec_name,
                emergency_contact_relation=ec_rel,
                emergency_contact_phone=ec_phone,
                is_phd=is_phd,
            ).on_conflict_do_nothing(constraint="uq_faculties_employee_id"),
        )

    return inserted


def _seed_faculty_simple_linear_process(
    session: Session,
    *,
    code: str,
    title: str,
    channel_role_codes: list[str],
) -> int:
    """Seed a faculty_* ApprovalProcess with Stage 1 resolver + linear channel (idempotent).

    Channel shape: dept_head_at_requestor_campus → remaining channel_role_codes.
    Used for the 4 Phase 5D linear-channel seeds (INVITED_TALK, PROFESSIONAL_MEMBERSHIP,
    WFH, FIELD_VISIT). Mirrors _seed_faculty_noc_process idempotency pattern.

    Returns total rows inserted (0–2).
    """
    from datetime import UTC, datetime

    from durgam.models.crosscutting import ApprovalProcess, ApprovalStageOption

    inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code=code,
            title=title,
            requestor_role_codes=["FACULTY"],
            channel_role_codes=channel_role_codes,
            is_finance=False,
            stage_pick_modes_json={"1": "approver"},
        )
        .on_conflict_do_nothing(constraint="uq_approval_processes_code"),
    )

    process = session.exec(
        select(ApprovalProcess).where(
            ApprovalProcess.code == code,
            ApprovalProcess.is_deleted == False,  # noqa: E712
        )
    ).first()

    if process is not None:
        # Idempotently set attachment config from defaults (preserves manual sys-admin edits).
        if process.max_upward_attachments == 0 and process.allowed_attachment_mime_types_json is None:
            process.max_upward_attachments = 3
            process.allowed_attachment_mime_types_json = ["application/pdf"]
            session.add(process)
            session.flush()

        existing_option = session.exec(
            select(ApprovalStageOption).where(
                ApprovalStageOption.approval_process_id == process.id,
                ApprovalStageOption.stage_index == 1,
                ApprovalStageOption.resolver_name == "dept_head_at_requestor_campus",
                ApprovalStageOption.is_deleted == False,  # noqa: E712
            )
        ).first()
        if existing_option is None:
            now = datetime.now(UTC)
            option = ApprovalStageOption(
                approval_process_id=process.id,
                stage_index=1,
                resolver_name="dept_head_at_requestor_campus",
                label="Head of Department",
                sort_order=0,
                created_at=now,
                updated_at=now,
            )
            session.add(option)
            session.flush()
            inserted += 1

    return inserted


def _seed_faculty_noc_process(session: Session) -> int:
    """Seed the faculty_noc ApprovalProcess and its Stage 1 OR-set option (idempotent).

    Stage 1: OR-set with pick_mode='approver' — pool is resolved by
    'dept_head_at_requestor_campus' (HoD→AhoD fallback for the requestor's dept+campus).
    Stage 2: legacy Registrar via channel_role_codes[1].

    Returns total rows inserted (0–2).
    """
    from datetime import UTC, datetime

    from durgam.models.crosscutting import ApprovalProcess, ApprovalStageOption

    inserted = _exec_insert(
        session,
        pg_insert(ApprovalProcess)
        .values(
            code="faculty_noc",
            title="Faculty No Objection Certificate",
            requestor_role_codes=["FACULTY"],
            channel_role_codes=["HOD", "REGISTRAR"],
            is_finance=False,
            stage_pick_modes_json={"1": "approver"},
        )
        .on_conflict_do_nothing(constraint="uq_approval_processes_code"),
    )

    process = session.exec(
        select(ApprovalProcess).where(
            ApprovalProcess.code == "faculty_noc",
            ApprovalProcess.is_deleted == False,  # noqa: E712
        )
    ).first()

    if process is not None:
        # Phase 6: idempotently upgrade attachment config from defaults.
        # Only update when values are still at defaults (preserves manual sys admin edits).
        if process.max_upward_attachments == 0 and process.allowed_attachment_mime_types_json is None:
            process.max_upward_attachments = 3
            process.allowed_attachment_mime_types_json = ["application/pdf"]
            session.add(process)
            session.flush()
        # Phase 7G: enable downward attachments for approver NOC response documents.
        if process.max_downward_attachments == 0 and not process.requires_downward_attachments:
            process.max_downward_attachments = 3
            session.add(process)
            session.flush()
        existing_option = session.exec(
            select(ApprovalStageOption).where(
                ApprovalStageOption.approval_process_id == process.id,
                ApprovalStageOption.stage_index == 1,
                ApprovalStageOption.resolver_name == "dept_head_at_requestor_campus",
                ApprovalStageOption.is_deleted == False,  # noqa: E712
            )
        ).first()
        if existing_option is None:
            now = datetime.now(UTC)
            option = ApprovalStageOption(
                approval_process_id=process.id,
                stage_index=1,
                resolver_name="dept_head_at_requestor_campus",
                label="Head of Department",
                sort_order=0,
                created_at=now,
                updated_at=now,
            )
            session.add(option)
            session.flush()
            inserted += 1

    return inserted


def main() -> None:
    engine = create_engine(settings.database_url_sync, echo=False)
    with Session(engine) as session:
        counts = seed(session)
    log.info("seed_complete", **counts)


if __name__ == "__main__":
    main()
