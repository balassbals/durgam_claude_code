"""Integration test for docgen — DOCX template-fill produces valid DOCX (E-005)."""

import io
from pathlib import Path

from docx import Document

from durgam.docgen.merge import render_docx_template


def _make_template_docx() -> bytes:
    doc = Document()
    doc.add_heading("{{ title }}", level=1)
    doc.add_paragraph("Date: {{ date }}")
    doc.add_paragraph("{{ body }}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocgenIntegration:
    def test_full_render_produces_valid_docx(self):
        tpl_bytes = _make_template_docx()
        context = {
            "title": "Board of Studies Meeting",
            "date": "2025-10-15",
            "body": "Minutes of the meeting held on 2025-10-15.",
        }
        result, warnings = render_docx_template(tpl_bytes, context)
        assert len(result) > 0
        assert warnings == []

        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Board of Studies Meeting" in all_text
        assert "2025-10-15" in all_text
        assert "Minutes of the meeting" in all_text

    def test_fixture_template_render(self):
        fixture_path = Path(__file__).resolve().parent.parent / "fixtures" / "sample_template.docx"
        if not fixture_path.exists():
            return
        tpl_bytes = fixture_path.read_bytes()
        result, warnings = render_docx_template(tpl_bytes, {
            "title": "Integration Test",
            "name": "Test User",
            "date": "2026-01-01",
        })
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Integration Test" in all_text
        assert "Test User" in all_text
