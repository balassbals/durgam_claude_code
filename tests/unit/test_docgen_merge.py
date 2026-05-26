"""Unit tests for docgen merge primitive — DOCX template-fill (E-005)."""

import io
from pathlib import Path

import pytest
from docx import Document

from durgam.docgen.merge import DocgenError, render_docx_template

_FIXTURE_DIR = Path(__file__).resolve().parent.parent / "fixtures"


def _make_template_docx(text_with_vars: str) -> bytes:
    doc = Document()
    doc.add_paragraph(text_with_vars)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestRenderDocxTemplate:
    def test_simple_variable_substitution(self):
        tpl = _make_template_docx("Hello {{ name }}!")
        result = render_docx_template(tpl, {"name": "World"})
        doc = Document(io.BytesIO(result))
        assert any("Hello World!" in p.text for p in doc.paragraphs)

    def test_empty_context_leaves_blank(self):
        tpl = _make_template_docx("Hello {{ name }}!")
        result = render_docx_template(tpl, {})
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs]
        assert any("Hello" in t for t in texts)

    def test_multiple_variables(self):
        tpl = _make_template_docx("{{ title }} by {{ author }}")
        result = render_docx_template(tpl, {"title": "Report", "author": "Admin"})
        doc = Document(io.BytesIO(result))
        assert any("Report by Admin" in p.text for p in doc.paragraphs)

    def test_invalid_docx_raises(self):
        with pytest.raises(DocgenError, match="rendering failed"):
            render_docx_template(b"not a docx", {"key": "val"})

    def test_returns_valid_docx_bytes(self):
        tpl = _make_template_docx("Test {{ x }}")
        result = render_docx_template(tpl, {"x": "value"})
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) >= 1

    def test_fixture_template(self):
        fixture_path = _FIXTURE_DIR / "sample_template.docx"
        if not fixture_path.exists():
            pytest.skip("sample_template.docx fixture not found")
        tpl_bytes = fixture_path.read_bytes()
        result = render_docx_template(tpl_bytes, {
            "title": "Test Report",
            "name": "John Doe",
            "date": "2026-05-26",
        })
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Test Report" in all_text
        assert "John Doe" in all_text
        assert "2026-05-26" in all_text
