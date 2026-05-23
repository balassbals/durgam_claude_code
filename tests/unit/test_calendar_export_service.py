"""Unit tests for CalendarExportService."""

import csv
import io
from datetime import UTC, datetime
from types import SimpleNamespace

import pytest

from durgam.services.calendar_export import CalendarExportService


def _entry(**overrides) -> SimpleNamespace:
    defaults = {
        "title": "Semester 1 Begins",
        "entry_type": "sem_begin",
        "starts_at": datetime(2025, 7, 1, 9, 0, tzinfo=UTC),
        "ends_at": datetime(2025, 7, 1, 10, 0, tzinfo=UTC),
        "owner_role_code": "REGISTRAR",
        "scope_type": None,
        "scope_id": None,
        "notes": None,
    }
    defaults.update(overrides)
    return SimpleNamespace(**defaults)


def _two_entries() -> list:
    return [
        _entry(),
        _entry(
            title="Annual Sports Day",
            entry_type="sports",
            starts_at=datetime(2025, 12, 15, 8, 0, tzinfo=UTC),
            ends_at=datetime(2025, 12, 15, 17, 0, tzinfo=UTC),
            owner_role_code="DIRECTOR",
            scope_type="campus",
            notes="All-day event",
        ),
    ]


class TestExportCSV:
    def test_empty_produces_header_only(self):
        svc = CalendarExportService()
        data = svc.export_csv([], "2025-26")
        reader = csv.reader(io.StringIO(data.decode("utf-8")))
        rows = list(reader)
        assert len(rows) == 1
        assert rows[0][0] == "Title"

    def test_two_entries_produce_three_rows(self):
        svc = CalendarExportService()
        data = svc.export_csv(_two_entries(), "2025-26")
        reader = csv.reader(io.StringIO(data.decode("utf-8")))
        rows = list(reader)
        assert len(rows) == 3
        assert rows[1][0] == "Semester 1 Begins"
        assert rows[2][0] == "Annual Sports Day"
        assert rows[2][5] == "campus"
        assert rows[2][6] == "All-day event"

    def test_returns_bytes(self):
        svc = CalendarExportService()
        result = svc.export_csv([], "2025-26")
        assert isinstance(result, bytes)


class TestExportExcel:
    def test_empty_produces_header_row(self):
        from openpyxl import load_workbook

        svc = CalendarExportService()
        data = svc.export_excel([], "2025-26")
        wb = load_workbook(io.BytesIO(data))
        ws = wb.active
        assert ws.cell(row=1, column=1).value == "Title"
        assert ws.max_row == 1

    def test_two_entries(self):
        from openpyxl import load_workbook

        svc = CalendarExportService()
        data = svc.export_excel(_two_entries(), "2025-26")
        wb = load_workbook(io.BytesIO(data))
        ws = wb.active
        assert ws.max_row == 3
        assert ws.cell(row=2, column=1).value == "Semester 1 Begins"
        assert ws.cell(row=3, column=1).value == "Annual Sports Day"

    def test_sheet_title_contains_ay_code(self):
        from openpyxl import load_workbook

        svc = CalendarExportService()
        data = svc.export_excel([], "2025-26")
        wb = load_workbook(io.BytesIO(data))
        assert "2025-26" in wb.active.title

    def test_returns_bytes(self):
        svc = CalendarExportService()
        result = svc.export_excel([], "2025-26")
        assert isinstance(result, bytes)


class TestExportPDF:
    def test_empty_produces_valid_pdf(self):
        svc = CalendarExportService()
        data = svc.export_pdf([], "2025-26")
        assert data[:5] == b"%PDF-"

    def test_two_entries_produces_valid_pdf(self):
        svc = CalendarExportService()
        data = svc.export_pdf(_two_entries(), "2025-26")
        assert data[:5] == b"%PDF-"
        assert len(data) > 100

    def test_returns_bytes(self):
        svc = CalendarExportService()
        result = svc.export_pdf([], "2025-26")
        assert isinstance(result, bytes)


class TestExportDOCX:
    def test_empty_produces_valid_docx(self):
        svc = CalendarExportService()
        data = svc.export_docx([], "2025-26")
        assert data[:2] == b"PK"

    def test_two_entries_produces_valid_docx(self):
        from docx import Document

        svc = CalendarExportService()
        data = svc.export_docx(_two_entries(), "2025-26")
        doc = Document(io.BytesIO(data))
        tables = doc.tables
        assert len(tables) == 1
        assert len(tables[0].rows) == 3

    def test_heading_contains_ay_code(self):
        from docx import Document

        svc = CalendarExportService()
        data = svc.export_docx([], "2025-26")
        doc = Document(io.BytesIO(data))
        heading = doc.paragraphs[0].text
        assert "2025-26" in heading

    def test_returns_bytes(self):
        svc = CalendarExportService()
        result = svc.export_docx([], "2025-26")
        assert isinstance(result, bytes)
