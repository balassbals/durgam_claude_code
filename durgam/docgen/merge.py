"""DOCX merge primitive — letterhead image + content blocks → DOCX bytes."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Inches


class DocgenError(Exception):
    pass


def merge_letterhead_and_content(
    letterhead_bytes: bytes,
    content_blocks: list[dict[str, Any]],
    *,
    mime_type: str = "image/png",
) -> bytes:
    """Create a DOCX with the letterhead in the header and content in the body.

    Args:
        letterhead_bytes: Raw image bytes (PNG or JPG only).
        content_blocks: List of dicts, each with a ``type`` key (heading,
            paragraph, table) and type-specific fields.
        mime_type: MIME type of the letterhead image.

    Returns:
        DOCX file bytes.

    Raises:
        DocgenError: If the letterhead is a PDF (not yet supported for merge).
    """
    if mime_type == "application/pdf":
        raise DocgenError(
            "PDF letterheads not yet supported for merge; use PNG or JPG."
        )

    doc = Document()

    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    run.add_picture(io.BytesIO(letterhead_bytes), width=Inches(6.0))

    for block in content_blocks:
        block_type = block.get("type", "")
        if block_type == "heading":
            doc.add_heading(block.get("text", ""), level=block.get("level", 1))
        elif block_type == "paragraph":
            doc.add_paragraph(block.get("text", ""))
        elif block_type == "table":
            rows_data = block.get("rows", [])
            if rows_data:
                cols = len(rows_data[0])
                table = doc.add_table(rows=len(rows_data), cols=cols)
                table.style = "Table Grid"
                for r_idx, row_data in enumerate(rows_data):
                    for c_idx, cell_val in enumerate(row_data):
                        table.rows[r_idx].cells[c_idx].text = str(cell_val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
