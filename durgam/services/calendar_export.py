"""CalendarExportService — CSV / Excel / PDF / DOCX export (§9.3 M4)."""

from __future__ import annotations

import csv
import io
from typing import TYPE_CHECKING

import structlog

if TYPE_CHECKING:
    from durgam.models.config_anchors import CalendarEntry

log = structlog.get_logger(__name__)

_COLUMNS = ("Title", "Start", "End", "Scope")


def _row(
    entry: CalendarEntry,
    scope_labels: dict[str, str] | None = None,
) -> tuple[str, ...]:
    scope = ""
    if entry.scope_type and entry.scope_id and scope_labels:
        scope = scope_labels.get(f"{entry.scope_type}:{entry.scope_id}", entry.scope_type)
    elif entry.scope_type:
        scope = entry.scope_type
    return (
        entry.title,
        entry.starts_at.strftime("%b %-d %I:%M %p"),
        entry.ends_at.strftime("%b %-d %I:%M %p"),
        scope,
    )


class CalendarExportService:
    def export_csv(
        self, entries: list[CalendarEntry], ay_code: str,
        scope_labels: dict[str, str] | None = None,
    ) -> bytes:
        buf = io.StringIO()
        writer = csv.writer(buf)
        writer.writerow(_COLUMNS)
        for entry in entries:
            writer.writerow(_row(entry, scope_labels))
        return buf.getvalue().encode("utf-8")

    def export_excel(
        self, entries: list[CalendarEntry], ay_code: str,
        scope_labels: dict[str, str] | None = None,
    ) -> bytes:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = f"Calendar {ay_code}"
        bold = Font(bold=True)
        for col_idx, col_name in enumerate(_COLUMNS, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.font = bold
        for row_idx, entry in enumerate(entries, 2):
            for col_idx, value in enumerate(_row(entry, scope_labels), 1):
                ws.cell(row=row_idx, column=col_idx, value=value)
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def export_pdf(
        self, entries: list[CalendarEntry], ay_code: str,
        scope_labels: dict[str, str] | None = None,
    ) -> bytes:
        from fpdf import FPDF

        _DEJAVU = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        _DEJAVU_B = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

        pdf = FPDF(orientation="L", format="A4")
        pdf.add_font("DejaVuSans", "", _DEJAVU)
        pdf.add_font("DejaVuSans", "B", _DEJAVU_B)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("DejaVuSans", "B", 14)
        pdf.cell(0, 10, f"Academic Calendar - {ay_code}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        pdf.set_font("DejaVuSans", "B", 9)
        col_widths = (80, 50, 50, 57)
        for i, col_name in enumerate(_COLUMNS):
            pdf.cell(col_widths[i], 8, col_name, border=1)
        pdf.ln()

        pdf.set_font("DejaVuSans", "", 8)
        for entry in entries:
            row = _row(entry, scope_labels)
            for i, value in enumerate(row):
                pdf.cell(col_widths[i], 7, value[:50], border=1)
            pdf.ln()

        return bytes(pdf.output())

    def export_docx(
        self, entries: list[CalendarEntry], ay_code: str,
        scope_labels: dict[str, str] | None = None,
    ) -> bytes:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(f"Academic Calendar - {ay_code}", level=1)

        table = doc.add_table(rows=1, cols=len(_COLUMNS))
        table.style = "Table Grid"
        for i, col_name in enumerate(_COLUMNS):
            cell = table.rows[0].cells[i]
            cell.text = col_name
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for entry in entries:
            row_cells = table.add_row().cells
            for i, value in enumerate(_row(entry, scope_labels)):
                row_cells[i].text = value
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
